"""API routes for apartments, documents, and file upload."""
import hashlib
import json
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, Body, File, Form, UploadFile, HTTPException
from pydantic import BaseModel

from database import get_connection

router = APIRouter(prefix="/api", tags=["api"])

UPLOAD_DIR = Path(__file__).parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Import unified LLM client
from llm_client import call_llm, is_llm_available, get_llm_info, LLM_MODEL, set_llm_mode, get_llm_mode, set_llm_model, get_llm_model, ping_llm, call_vision_llm

# Import RAG index (FAISS + embedding)
import numpy as np
from rag_index import get_embedding, build_index as rag_build_index, search as faiss_search, invalidate_index

@router.get("/debug/llm")
def debug_llm():
    """Debug endpoint to check LLM config."""
    import os
    return {
        "is_llm_available": is_llm_available(),
        "get_llm_info": get_llm_info(),
        "env_LLM_ENABLED": os.getenv("LLM_ENABLED"),
        "env_LLM_PROVIDER": os.getenv("LLM_PROVIDER"),
        "env_LLM_API_KEY_set": bool(os.getenv("LLM_API_KEY"))
    }


@router.get("/llm/mode")
def get_mode():
    """Get current LLM mode (local/remote)."""
    return {"mode": get_llm_mode(), "info": get_llm_info()}


def _fetch_model_list() -> list[dict]:
    """Fetch available models from current Ollama endpoint."""
    import requests as _req
    info = get_llm_info()
    base_url = (info.get("base_url") or "http://127.0.0.1:11434").rstrip("/")
    headers = {}
    api_key = os.getenv("LLM_API_KEY", "")
    if info.get("mode") == "remote" and api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    try:
        resp = _req.get(f"{base_url}/api/tags", headers=headers, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        models = []
        for m in data.get("models", []):
            name = m.get("name", "")
            size_gb = round(m.get("size", 0) / 1e9, 1) if m.get("size") else None
            models.append({"name": name, "size_gb": size_gb})
        models.sort(key=lambda x: x["name"])
        return models
    except Exception:
        return []


@router.post("/llm/mode")
def switch_mode(mode: str = Form(...)):
    """Switch LLM mode at runtime. Returns model list for the new mode."""
    try:
        set_llm_mode(mode)
        models = _fetch_model_list()
        return {"success": True, "mode": get_llm_mode(), "info": get_llm_info(), "models": models}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/llm/model")
def get_model():
    """Get current LLM model."""
    return {"model": get_llm_model(), "info": get_llm_info()}


@router.post("/llm/model")
def switch_model(model: str = Form(...)):
    """Switch LLM model at runtime."""
    set_llm_model(model)
    return {"success": True, "model": get_llm_model(), "info": get_llm_info()}


@router.get("/llm/models")
def list_models():
    """List available models from current Ollama endpoint (local or remote)."""
    info = get_llm_info()
    models = _fetch_model_list()
    return {"models": models, "mode": info.get("mode"), "current": get_llm_model()}


class ApartmentCreate(BaseModel):
    apt_id: str
    name: str


class ApartmentResponse(BaseModel):
    apt_id: str
    name: str
    created_at: str


# ============ APARTMENTS ============

@router.post("/apartments")
def create_apartment(apt_id: str = Form(...), name: str = Form(...)):
    """Create a new apartment (accepts form data)."""
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute(
            "INSERT INTO apartments (apt_id, name) VALUES (%s, %s)",
            (apt_id, name)
        )
        conn.commit()
        return {"success": True, "apt_id": apt_id}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
    finally:
        conn.close()


@router.get("/apartments")
def list_apartments():
    """List all apartments."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT apt_id, name, created_at FROM apartments ORDER BY created_at DESC")
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


@router.delete("/apartments/{apt_id}")
def delete_apartment(apt_id: str):
    """Delete apartment and all related data (cascading)."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT apt_id FROM apartments WHERE apt_id = %s", (apt_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Apartment not found")

    cursor.execute("SELECT doc_id FROM documents WHERE apt_id = %s", (apt_id,))
    doc_ids = [r["doc_id"] for r in cursor.fetchall()]
    cursor.execute("SELECT conversation_id FROM conversations WHERE apt_id = %s", (apt_id,))
    conv_ids = [r["conversation_id"] for r in cursor.fetchall()]
    deleted = {}
    if doc_ids:
        ph = ",".join(["%s"] * len(doc_ids))
        for tbl in ["manual_section_revisions", "manual_sections", "qa_issues", "chunks", "api_specs", "doc_chunks"]:
            cursor.execute(f"DELETE FROM {tbl} WHERE doc_id IN ({ph})", doc_ids)
            deleted[tbl] = cursor.rowcount
    if conv_ids:
        ph2 = ",".join(["%s"] * len(conv_ids))
        cursor.execute(f"DELETE FROM messages WHERE conversation_id IN ({ph2})", conv_ids)
        deleted["messages"] = cursor.rowcount
    for tbl in ["conversations", "improve_suggestions", "branch_class_cache", "documents"]:
        cursor.execute(f"DELETE FROM {tbl} WHERE apt_id = %s", (apt_id,))
        deleted[tbl] = cursor.rowcount
    cursor.execute("DELETE FROM apartments WHERE apt_id = %s", (apt_id,))
    deleted["apartments"] = cursor.rowcount
    conn.commit()
    conn.close()
    return {"success": True, "apt_id": apt_id, "deleted": deleted}


@router.delete("/doc/{doc_id}")
def delete_document(doc_id: str):
    """Delete a document and all related data."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT doc_id FROM documents WHERE doc_id = %s", (doc_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")
    deleted = {}
    for tbl in ["manual_section_revisions", "manual_sections", "qa_issues", "chunks", "api_specs", "doc_chunks"]:
        cursor.execute(f"DELETE FROM {tbl} WHERE doc_id = %s", (doc_id,))
        deleted[tbl] = cursor.rowcount
    cursor.execute("DELETE FROM documents WHERE doc_id = %s", (doc_id,))
    deleted["documents"] = cursor.rowcount
    conn.commit()
    conn.close()
    return {"success": True, "doc_id": doc_id, "deleted": deleted}


# ============ UPLOAD ============

@router.post("/upload")
async def upload_document(
    apt_id: str = Form(...),
    file: UploadFile = File(...)
):
    """Upload a document file."""
    content = await file.read()
    content_hash = hashlib.sha256(content).hexdigest()
    
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    source_type = ext if ext in ("docx", "pdf", "txt", "md") else "txt"
    
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute(
        "SELECT doc_id, version FROM documents WHERE apt_id = %s AND content_hash = %s AND status != 'ARCHIVED' ORDER BY version DESC LIMIT 1",
        (apt_id, content_hash)
    )
    existing = cursor.fetchone()
    
    if existing:
        new_version = existing["version"] + 1
        cursor.execute(
            "UPDATE documents SET status = 'ARCHIVED', updated_at = %s WHERE apt_id = %s AND content_hash = %s AND status != 'ARCHIVED'",
            (datetime.now().isoformat(), apt_id, content_hash)
        )
    else:
        cursor.execute("SELECT MAX(version) as max_ver FROM documents WHERE apt_id = %s", (apt_id,))
        row = cursor.fetchone()
        new_version = (row["max_ver"] or 0) + 1
    
    doc_id = f"doc_{uuid.uuid4().hex[:12]}"
    file_path = UPLOAD_DIR / f"{doc_id}.{source_type}"
    with open(file_path, "wb") as f:
        f.write(content)
    
    raw_text = f"[Placeholder: extract text first]"
    
    now = datetime.now().isoformat()
    cursor.execute("""
        INSERT INTO documents (doc_id, apt_id, title, source_filename, source_type, content_hash, raw_text, version, status, created_at, updated_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'DRAFT', %s, %s)
    """, (doc_id, apt_id, filename, filename, source_type, content_hash, raw_text, new_version, now, now))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "doc_id": doc_id, "version": new_version, "content_hash": content_hash, "status": "DRAFT"}


@router.get("/docs")
def list_documents(apt_id: Optional[str] = None):
    """List documents, optionally filtered by apartment."""
    conn = get_connection()
    cursor = conn.cursor()
    
    if apt_id:
        cursor.execute(
            "SELECT doc_id, apt_id, title, source_filename, source_type, version, status, created_at FROM documents WHERE apt_id = %s ORDER BY created_at DESC",
            (apt_id,)
        )
    else:
        cursor.execute(
            "SELECT doc_id, apt_id, title, source_filename, source_type, version, status, created_at FROM documents ORDER BY created_at DESC"
        )
    
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]


# ============ STEP 2: EXTRACT TEXT ============

VISION_PROMPT = "Ïù¥ Ïù¥ÎØ∏ÏßÄÎ•º ÌïúÍµ≠Ïñ¥Î°ú ÏûêÏÑ∏Ìûà ÏÑ§Î™ÖÌï¥Ï§ò. UI ÌôîÎ©¥Ïù¥ÎùºÎ©¥ Ïñ¥Îñ§ Í∏∞Îä•Ïùò ÌôîÎ©¥Ïù∏ÏßÄ, Î≤ÑÌäº/Î©îÎâ¥/ÏûÖÎ†• ÌïÑÎìú Îì± Íµ¨ÏÑ± ÏöîÏÜåÎ•º Ìè¨Ìï®Ìï¥ÏÑú ÏÑ§Î™ÖÌï¥."

# Extract progress tracker (in-memory)
_extract_progress = {}  # {doc_id: {"page": 1, "total_pages": 10, "images_done": 2, "images_total": 5, "status": "..."}}


def _is_extract_cancelled(doc_id: str) -> bool:
    return _extract_progress.get(doc_id, {}).get("cancelled", False)


def _describe_image(image_base64: str, doc_id: str = None) -> str:
    """Vision LLMÏúºÎ°ú Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö ÏÉùÏÑ±. Ïã§Ìå® Ïãú [Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Ïã§Ìå®] Î∞òÌôò."""
    if doc_id and _is_extract_cancelled(doc_id):
        return "[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Í±¥ÎÑàÎúÄ: Ï∑®ÏÜåÎê®]"
    if doc_id and doc_id in _extract_progress:
        _extract_progress[doc_id]["images_done"] = _extract_progress[doc_id].get("images_done", 0) + 1
        _extract_progress[doc_id]["status"] = f"Ïù¥ÎØ∏ÏßÄ {_extract_progress[doc_id]['images_done']}/{_extract_progress[doc_id].get('images_total', '?')} Vision LLM Î∂ÑÏÑù Ï§ë"
    try:
        desc = call_vision_llm(VISION_PROMPT, image_base64)
        if desc and desc.strip():
            return f"[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö: {desc.strip()}]"
    except Exception as e:
        print(f"[EXTRACT] Vision LLM failed: {e}")
    return "[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Ïã§Ìå®]"


@router.get("/doc/{doc_id}/extract-progress")
def get_extract_progress(doc_id: str):
    """Poll extract progress."""
    return _extract_progress.get(doc_id, {})


@router.post("/doc/{doc_id}/extract-cancel")
def cancel_extract(doc_id: str):
    """Cancel an in-progress extract operation."""
    if doc_id in _extract_progress:
        _extract_progress[doc_id]["cancelled"] = True
        return {"success": True, "detail": "Ï∑®ÏÜå ÏöîÏ≤≠Îê®"}
    return {"success": False, "detail": "ÏßÑÌñâ Ï§ëÏù∏ ÏûëÏóÖ ÏóÜÏùå"}


@router.post("/doc/{doc_id}/extract-text")
def extract_text(doc_id: str, resume_page: int = 0):
    """Extract text from uploaded document (DOCX, PDF, TXT, MD).
    resume_page: 0=Ï≤òÏùåÎ∂ÄÌÑ∞, N=NÌéòÏù¥ÏßÄÎ∂ÄÌÑ∞ Ïù¥Ïñ¥ÏÑú (PDF only, Í∏∞Ï°¥ raw_textÏóê append)
    """
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT source_filename, source_type, raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")

    source_type = doc["source_type"]
    file_path = UPLOAD_DIR / f"{doc_id}.{source_type}"

    if not file_path.exists():
        conn.close()
        raise HTTPException(status_code=404, detail="File not found on disk")

    # Ïù¥Ïñ¥ÌïòÍ∏∞: Í∏∞Ï°¥ ÌÖçÏä§Ìä∏ Î≥¥Ï°¥
    existing_text = (doc["raw_text"] or "") if resume_page > 0 else ""
    raw_text = ""
    image_count = 0

    if source_type == "docx":
        try:
            from docx import Document
            import zipfile
            import base64

            _extract_progress[doc_id] = {"page": 0, "total_pages": 0, "images_done": 0, "images_total": 0, "status": "DOCX ÌÖçÏä§Ìä∏ Ï∂îÏ∂ú Ï§ë", "pages": {}}

            docx_doc = Document(str(file_path))
            parts = []

            # Îã®ÎùΩ
            para_texts = []
            for p in docx_doc.paragraphs:
                if p.text.strip():
                    para_texts.append(p.text)
            if para_texts:
                para_content = "\n".join(para_texts)
                parts.append(para_content)
                _extract_progress[doc_id]["pages"]["Î≥∏Î¨∏"] = para_content

            # ÌÖåÏù¥Î∏î
            _extract_progress[doc_id]["status"] = "ÌÖåÏù¥Î∏î Ï∂îÏ∂ú Ï§ë"
            for t_idx, table in enumerate(docx_doc.tables):
                rows_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_text.append(" | ".join(cells))
                if rows_text:
                    table_content = "[Ìëú]\n" + "\n".join(rows_text)
                    parts.append(table_content)
                    _extract_progress[doc_id]["pages"][f"Ìëú {t_idx + 1}"] = table_content

            # Ïù¥ÎØ∏ÏßÄ Ïàò ÎØ∏Î¶¨ Ïπ¥Ïö¥Ìä∏
            with zipfile.ZipFile(str(file_path), 'r') as zf:
                media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
                _extract_progress[doc_id]["images_total"] = len(media_files)
                _extract_progress[doc_id]["status"] = f"Ïù¥ÎØ∏ÏßÄ 0/{len(media_files)} Vision LLM Î∂ÑÏÑù Ï§ë" if media_files else "ÏôÑÎ£å Ï§ë"

                for name in media_files:
                    img_bytes = zf.read(name)
                    img_b64 = base64.b64encode(img_bytes).decode()
                    image_count += 1
                    print(f"[EXTRACT] DOCX image {image_count}: {name} ({len(img_bytes)} bytes)")
                    desc = _describe_image(img_b64, doc_id)
                    parts.append(desc)
                    _extract_progress[doc_id]["pages"][f"Ïù¥ÎØ∏ÏßÄ {image_count}"] = desc

            raw_text = "\n".join(parts)
            if _is_extract_cancelled(doc_id):
                _extract_progress[doc_id].pop("pages", None)
                _extract_progress[doc_id]["status"] = "Ï∑®ÏÜåÎê®"
            else:
                _extract_progress.pop(doc_id, None)
        except Exception as e:
            _extract_progress.pop(doc_id, None)
            conn.close()
            raise HTTPException(status_code=500, detail=f"DOCX extraction failed: {str(e)}")

    elif source_type == "pdf":
        try:
            import fitz
            import base64

            pdf_doc = fitz.open(str(file_path))
            total_pages = len(pdf_doc)

            # Ïù¥ÎØ∏ÏßÄ Ìè¨Ìï® ÌéòÏù¥ÏßÄ Ïàò ÎØ∏Î¶¨ Ïπ¥Ïö¥Ìä∏
            pages_with_images = []
            for pg_idx, pg in enumerate(pdf_doc):
                if pg.get_images(full=True):
                    pages_with_images.append(pg_idx)

            # resume_page: 0-indexed internally (resume_page=3 ‚Üí skip pages 0,1,2)
            start_page = max(0, resume_page - 1) if resume_page > 0 else 0
            remaining_image_pages = [p for p in pages_with_images if p >= start_page]
            _extract_progress[doc_id] = {"page": start_page, "total_pages": total_pages, "images_done": 0, "images_total": len(remaining_image_pages), "status": f"{start_page}/{total_pages} ÌéòÏù¥ÏßÄ", "pages": {}}

            if start_page > 0:
                print(f"[EXTRACT] PDF resume from page {start_page + 1}/{total_pages} for {doc_id}")

            # Phase 1: ÌÖçÏä§Ìä∏ Ï∂îÏ∂ú + Ïù¥ÎØ∏ÏßÄ ÌéòÏù¥ÏßÄ Î†åÎçîÎßÅ (ÏàúÏ∞®, Îπ†Î¶Ñ)
            page_texts = {}       # {page_num: text}
            image_renders = {}    # {page_num: img_b64}
            for page_num, page in enumerate(pdf_doc):
                if page_num < start_page:
                    continue
                _extract_progress[doc_id]["page"] = page_num + 1
                _extract_progress[doc_id]["status"] = f"{page_num + 1}/{total_pages} ÌéòÏù¥ÏßÄ ÌÖçÏä§Ìä∏ Ï∂îÏ∂ú"

                page_text = page.get_text().strip()
                if page_text:
                    page_texts[page_num] = page_text
                    # ÌÖçÏä§Ìä∏Îßå ÏûàÎäî ÌéòÏù¥ÏßÄÎäî Î∞îÎ°ú Ïä§Ìä∏Î¶¨Î∞ç
                    if page_num not in pages_with_images:
                        _extract_progress[doc_id]["pages"][str(page_num + 1)] = page_text

                if page_num in remaining_image_pages:
                    pix = page.get_pixmap(dpi=150)
                    image_renders[page_num] = base64.b64encode(pix.tobytes("png")).decode()
                    print(f"[EXTRACT] PDF page {page_num + 1} render: {pix.width}x{pix.height}")

            pdf_doc.close()

            # Phase 2: Vision LLM Î≥ëÎ†¨ Ìò∏Ï∂ú (Î≥ëÎ™© Íµ¨Í∞Ñ)
            from concurrent.futures import ThreadPoolExecutor, as_completed
            VISION_WORKERS = min(3, len(image_renders)) or 1
            image_descs = {}  # {page_num: desc}

            if image_renders:
                _extract_progress[doc_id]["status"] = f"Ïù¥ÎØ∏ÏßÄ {len(image_renders)}ÌéòÏù¥ÏßÄ Vision LLM Î≥ëÎ†¨ Î∂ÑÏÑù Ï§ë"
                print(f"[EXTRACT] Vision LLM parallel: {len(image_renders)} pages, {VISION_WORKERS} workers")

                def _vision_task(pg_num, img_b64):
                    if _is_extract_cancelled(doc_id):
                        return pg_num, "[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Í±¥ÎÑàÎúÄ: Ï∑®ÏÜåÎê®]"
                    try:
                        desc = call_vision_llm(VISION_PROMPT, img_b64)
                        return pg_num, f"[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö: {desc.strip()}]" if desc and desc.strip() else (pg_num, "[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Ïã§Ìå®]")
                    except Exception as e:
                        print(f"[EXTRACT] Vision LLM failed page {pg_num + 1}: {e}")
                        return pg_num, "[Ïù¥ÎØ∏ÏßÄ ÏÑ§Î™Ö Ïã§Ìå®]"

                with ThreadPoolExecutor(max_workers=VISION_WORKERS) as executor:
                    futures = {executor.submit(_vision_task, pn, b64): pn for pn, b64 in image_renders.items()}
                    for future in as_completed(futures):
                        pg_num, desc = future.result()
                        image_descs[pg_num] = desc
                        image_count += 1
                        _extract_progress[doc_id]["images_done"] = image_count
                        _extract_progress[doc_id]["status"] = f"Ïù¥ÎØ∏ÏßÄ Î∂ÑÏÑù {image_count}/{len(image_renders)}"
                        # Ìï¥Îãπ ÌéòÏù¥ÏßÄ ÌÖçÏä§Ìä∏+Ïù¥ÎØ∏ÏßÄ Ìï©Ï≥êÏÑú Ïä§Ìä∏Î¶¨Î∞ç
                        combined = page_texts.get(pg_num, "")
                        combined = (combined + "\n" + desc).strip() if combined else desc
                        _extract_progress[doc_id]["pages"][str(pg_num + 1)] = combined

            # Phase 3: ÌéòÏù¥ÏßÄ ÏàúÏÑúÎåÄÎ°ú Ï°∞Ìï©
            parts = []
            all_pages = sorted(set(list(page_texts.keys()) + list(image_descs.keys())))
            for pg_num in all_pages:
                page_parts = []
                if pg_num in page_texts:
                    page_parts.append(page_texts[pg_num])
                if pg_num in image_descs:
                    page_parts.append(image_descs[pg_num])
                if page_parts:
                    parts.append("\n".join(page_parts))

            pdf_doc.close()
            raw_text = "\n".join(parts)
            if _is_extract_cancelled(doc_id):
                # Ï∑®ÏÜå Ïãú: progress Ïú†ÏßÄ (Ïù¥Ïñ¥ÌïòÍ∏∞Ïö©), pagesÎßå Ï†úÍ±∞ (Î©îÎ™®Î¶¨ Ï†àÏïΩ)
                _extract_progress[doc_id].pop("pages", None)
                _extract_progress[doc_id]["status"] = f"Ï∑®ÏÜåÎê® ({_extract_progress[doc_id].get('page', 0)}/{total_pages} ÌéòÏù¥ÏßÄ)"
            else:
                _extract_progress.pop(doc_id, None)
        except Exception as e:
            # ÏóêÎü¨ ÏãúÏóêÎèÑ progress Ïú†ÏßÄ (Ïù¥Ïñ¥ÌïòÍ∏∞Ïö©)
            if doc_id in _extract_progress:
                _extract_progress[doc_id].pop("pages", None)
                _extract_progress[doc_id]["status"] = f"Ïò§Î•ò ({_extract_progress[doc_id].get('page', 0)}/{_extract_progress[doc_id].get('total_pages', '?')} ÌéòÏù¥ÏßÄ)"
            conn.close()
            raise HTTPException(status_code=500, detail=f"PDF extraction failed: {str(e)}")

    elif source_type in ("txt", "md"):
        raw_text = file_path.read_text(encoding="utf-8", errors="ignore")

    else:
        raw_text = f"[ÏßÄÏõêÌïòÏßÄ ÏïäÎäî ÌååÏùº ÌòïÏãù: {source_type}]"

    # Ïù¥Ïñ¥ÌïòÍ∏∞: Í∏∞Ï°¥ ÌÖçÏä§Ìä∏ + ÏÉà ÌÖçÏä§Ìä∏ Ìï©Ïπ®
    if existing_text and raw_text:
        raw_text = existing_text + "\n" + raw_text
    elif existing_text:
        raw_text = existing_text

    cursor.execute("UPDATE documents SET raw_text = %s, updated_at = %s WHERE doc_id = %s",
                   (raw_text, datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()

    return {
        "success": True,
        "doc_id": doc_id,
        "chars": len(raw_text),
        "resumed_from": resume_page if resume_page > 0 else None,
        "image_count": image_count,
        "preview": raw_text[:300] if raw_text else ""
    }


MANUALIZE_HARD_LIMIT = 30000  # chars ‚Äî force window split above this (128K ctx Í∏∞Ï§Ä)
MANUALIZE_WINDOW_SIZE = 25000  # chars per window (~37K tokens, 128K ctxÏùò ~75%)
MANUALIZE_WINDOW_OVERLAP = 300

# Manualize progress tracker (in-memory, per doc_id)
_manualize_progress = {}  # {doc_id: {"done": 3, "total": 15}}


@router.get("/doc/{doc_id}/char-count")
def get_char_count(doc_id: str):
    """Get raw_text character count and window split info (lightweight pre-check)."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    raw_text = row["raw_text"] or ""
    chars = len(raw_text)
    mode = get_llm_mode()
    window_count = len(_split_by_headings(raw_text))
    return {"chars": chars, "window_count": window_count, "hard_limit": MANUALIZE_HARD_LIMIT, "mode": mode}


@router.get("/doc/{doc_id}/manualize-progress")
def get_manualize_progress(doc_id: str):
    """Poll manualize progress. Returns {done, total, cancelled, sections}."""
    prog = _manualize_progress.get(doc_id, {"done": 0, "total": 0})
    return {
        "done": prog.get("done", 0),
        "total": prog.get("total", 0),
        "cancelled": prog.get("cancelled", False),
        "sections": prog.get("sections", {}),
    }


@router.post("/doc/{doc_id}/manualize-cancel")
def cancel_manualize(doc_id: str):
    """Cancel an in-progress manualize operation."""
    if doc_id in _manualize_progress:
        _manualize_progress[doc_id]["cancelled"] = True
        return {"success": True, "detail": "Ï∑®ÏÜå ÏöîÏ≤≠Îê®"}
    return {"success": False, "detail": "ÏßÑÌñâ Ï§ëÏù∏ ÏûëÏóÖ ÏóÜÏùå"}


# ============ STEP 2: MANUALIZE ============

MANUALIZE_PROMPT = """ÎãπÏã†ÏùÄ ÏòÅÏóÖ/Ïö¥ÏòÅ Î¨∏ÏÑúÎ•º RAG(Í≤ÄÏÉâ Í∏∞Î∞ò ÎãµÎ≥Ä)Ïóê ÎÑ£Í∏∞ ÏúÑÌïú
"Íµ¨Ï°∞ÌôîÎêú Ï†ïÎ≥¥ Ï∂îÏ∂úÍ∏∞"ÏûÖÎãàÎã§.

‚ö†Ô∏è Ïù¥ ÏûëÏóÖÏùÄ ÏöîÏïΩÏù¥ ÏïÑÎãôÎãàÎã§.
‚ö†Ô∏è Î¨∏ÏÑúÎ•º Ï§ÑÏù¥Í±∞ÎÇò ÏïïÏ∂ïÌïòÎäî ÏûëÏóÖÏù¥ ÏïÑÎãôÎãàÎã§.

[ÏµúÏö∞ÏÑ† Î™©Ìëú]
- ÏõêÎ¨∏(raw_text)Ïùò Ï†ïÎ≥¥ÏôÄ Íµ¨Ï°∞Î•º ÏµúÎåÄÌïú Í∑∏ÎåÄÎ°ú Î≥¥Ï°¥ÌïòÏó¨ JSONÏúºÎ°ú Íµ¨Ï°∞ÌôîÌï©ÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏóÜÎäî Ï†ïÎ≥¥Î•º Ï∂îÍ∞Ä/Ï∂îÏ∏°/Ï∞ΩÏûëÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏûàÎäî Ï†ïÎ≥¥Î•º ÏÇ≠Ï†úÌïòÍ±∞ÎÇò ÏÉùÎûµÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- Î¨∏ÏÑúÎ•º Í∞ÑÎûµÌôîÌïòÍ±∞ÎÇò ÏïïÏ∂ïÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- Í∞ÄÎä•Ìïú Ìïú ÏõêÎ¨∏ Ï†ïÎ≥¥ÎüâÏùÑ Ïú†ÏßÄÌï©ÎãàÎã§.
- RAG Í≤ÄÏÉâÏóê Ïûò Í±∏Î¶¨ÎèÑÎ°ù Íµ¨Ï°∞Îßå Ï†ïÎ¶¨Ìï©ÎãàÎã§.

Ï¶â:
ÏöîÏïΩ ‚ùå
Ïû¨ÏûëÏÑ± ‚ùå
Íµ¨Ï°∞ÌôîÎêú Ï∂îÏ∂ú ‚úÖ

[ÏûÖÎ†•]
raw_text: {raw_text}

[Ï∂úÎ†• ÌòïÏãù]
- RFC8259 Ïú†Ìö® JSONÎßå Ï∂úÎ†•
- ÏΩîÎìúÎ∏îÎ°ù, ÏÑ§Î™Ö, Ï£ºÏÑù, ÎßàÌÅ¨Îã§Ïö¥ Í∏àÏßÄ
- Ïò§ÏßÅ JSON ÌÖçÏä§Ìä∏Îßå
- Î™®Îì† Î¨∏ÏûêÏó¥ÏùÄ ÌÅ∞Îî∞Ïò¥Ìëú ÏÇ¨Ïö©
- trailing comma Í∏àÏßÄ
- Î™®Îì† ÌïÑÎìúÎäî Î∞òÎìúÏãú Ìè¨Ìï® (ÏóÜÏúºÎ©¥ "" ÎòêÎäî [])

{{
  "doc_title": "",
  "doc_type": "POLICY|PROCESS|FAQ|NOTICE|MIXED",
  "summary": "",
  "sections": [
    {{
      "section_id": "",
      "name": "",
      "tags": [],
      "content": [
        {{
          "rule_id": "",
          "title": "",
          "bullets": [],
          "structured": {{
            "target": "",
            "condition": "",
            "procedure": [],
            "exceptions": [],
            "owner": "",
            "channel": ""
          }},
          "source_quotes": [],
          "issues": []
        }}
      ]
    }}
  ],
  "clarification_questions": [],
  "pii_handling": {{
    "pii_found": false,
    "pii_types": [],
    "masking_policy": []
  }},
  "change_summary": ""
}}

[üö® Ï†àÎåÄ Í∑úÏπô (ÏúÑÎ∞ò Í∏àÏßÄ)]

1) Ï†ïÎ≥¥ ÏÇ≠Ï†ú Í∏àÏßÄ
- ÏõêÎ¨∏Ïóê ÏûàÎäî Ìï≠Î™©/Î¶¨Ïä§Ìä∏/Î¨∏Ïû•ÏùÑ ÏûÑÏùòÎ°ú Ï†úÍ±∞ÌïòÏßÄ ÎßàÏÑ∏Ïöî.

2) Ï†ïÎ≥¥ ÏïïÏ∂ï Í∏àÏßÄ
- Ïó¨Îü¨ Ìï≠Î™©ÏùÑ ÌïòÎÇòÎ°ú Ìï©Ï≥ê Ï§ÑÏù¥ÏßÄ ÎßàÏÑ∏Ïöî.

3) ÏöîÏïΩ Í∏àÏßÄ
- ÏùòÎØ∏Îßå ÎÇ®Í∏∞Í≥† Ï∂ïÏïΩÌïòÏßÄ ÎßàÏÑ∏Ïöî.

4) ÏùºÎ∞òÌôî Í∏àÏßÄ
- "Îì±", "Í∏∞ÌÉÄ", "Ìè¨Ìï®"ÏúºÎ°ú Î≠âÍ∞úÏßÄ ÎßàÏÑ∏Ïöî.

5) Ïû¨Íµ¨ÏÑ± ÏµúÏÜåÌôî
- Íµ¨Ï°∞ Ï†ïÎ¶¨ Ïô∏ ÏùòÎØ∏ Î≥ÄÍ≤Ω Í∏àÏßÄ.

[ÏÑπÏÖò Í∑úÏπô - ÏµúÏö∞ÏÑ†]

- sectionsÎäî ÏõêÎ¨∏ Ìó§Îî© Íµ¨Ï°∞Î•º Í∑∏ÎåÄÎ°ú Îî∞Î¶ÖÎãàÎã§.
- ÏÑπÏÖò ÏàòÎ•º ÏûÑÏùòÎ°ú ÎäòÎ¶¨Í±∞ÎÇò Ï§ÑÏù¥ÏßÄ ÎßàÏÑ∏Ïöî.
- Î≥ëÌï©/Î∂ÑÎ¶¨ Í∏àÏßÄ.
- ÏõêÎ¨∏ Ìó§Îî©Ïù¥ ÏóÜÏúºÎ©¥ sectionsÎäî 1Í∞ú("general").

section_id:
- ÏÑπÏÖòÎ™Ö slug ÏÇ¨Ïö©
- ÏóÜÏúºÎ©¥ "general"

[rule ÏÉùÏÑ± Í∑úÏπô]

- ÏõêÎ¨∏ Î∂àÎ¶ø Î¶¨Ïä§Ìä∏Îäî Í∞úÏàòÏôÄ Ìï≠Î™©ÏùÑ Í∑∏ÎåÄÎ°ú Î≥¥Ï°¥.
- Ìïú Î∂àÎ¶ø = Ìïú Ï†ïÎ≥¥ Îã®ÏúÑ.
- ÏÇ≠Ï†ú/Î≥ëÌï© Í∏àÏßÄ.
- "ÌòÑÏû¨ Î≤ÑÏ†Ñ", "ÏõπÏÇ¨Ïù¥Ìä∏" Í∞ôÏùÄ Ï†ïÎ≥¥ÎèÑ Í∞ÅÍ∞Å Î≥ÑÎèÑ rule Í∞ÄÎä•.

rule_id:
S1-R1, S1-R2‚Ä¶ ÏàúÏÑú Í≥†Ï†ï

[bullets ÏûëÏÑ± Í∑úÏπô]

- ÏõêÎ¨∏ Î¨∏Ïû•ÏùÑ ÏµúÎåÄÌïú Ïú†ÏßÄ
- Í≥ºÎèÑÌïú Ïû¨ÏûëÏÑ± Í∏àÏßÄ
- Ï†ïÎ≥¥ Ï∂îÍ∞Ä Í∏àÏßÄ

[source_quotes]

- Í∞Å ruleÎßàÎã§ 0~2Í∞ú
- ÏõêÎ¨∏ÏóêÏÑú Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨
- ÌïúÍµ≠Ïñ¥ 20~70Ïûê
- ÏóÜÏúºÎ©¥ []

[structured]

- ÏõêÎ¨∏Ïóê Î™ÖÏãúÎêú Í≤ÉÎßå Ï±ÑÏõÄ
- ÏóÜÏúºÎ©¥ Ï†ÑÎ∂Ä ÎπàÍ∞í

[issues]

- Î¨∏Ï†ú ÏóÜÏúºÎ©¥ []

[PII]

ÏûàÏúºÎ©¥:
- ÎßàÏä§ÌÇπ
- issuesÏóê PII_RISK
- pii_found=true

ÏóÜÏúºÎ©¥:
- pii_found=false
- ÎÇòÎ®∏ÏßÄ []

[doc_type ÌåêÎã®]

POLICY: Í∑úÏ†ï/Í∏∞Ï§Ä
PROCESS: Ï†àÏ∞®
FAQ: Q/A
NOTICE: ÏïàÎÇ¥/ÏÜåÍ∞ú
MIXED: ÌòºÌï©

[ÎßàÏßÄÎßâ Ï≤¥ÌÅ¨ (Ïä§Ïä§Î°ú Í≤ÄÏ¶ù)]

JSON Ï∂úÎ†• Ï†Ñ Î∞òÎìúÏãú ÌôïÏù∏:
- ÏõêÎ¨∏ Ï†ïÎ≥¥Í∞Ä Îπ†ÏßÄÏßÄ ÏïäÏïòÎäîÍ∞Ä?
- Î∂àÎ¶ø Í∞úÏàòÍ∞Ä Ï§ÑÏßÄ ÏïäÏïòÎäîÍ∞Ä?
- ÏöîÏïΩÌïòÏßÄ ÏïäÏïòÎäîÍ∞Ä?

ÌïòÎÇòÎùºÎèÑ ÏúÑÎ∞òÏù¥Î©¥ Îã§Ïãú ÏÉùÏÑ±ÌïòÏÑ∏Ïöî.

[Í∏àÏßÄ]
- Ï∂îÏ∏°
- ÏùºÎ∞ò ÏÉÅÏãù Î≥¥ÏôÑ
- ÏÉà Ï†ïÎ≥¥ Ï∂îÍ∞Ä
- JSON Ïô∏ Ï∂úÎ†•"""


@router.post("/doc/{doc_id}/manualize")
def manualize(doc_id: str, force: bool = False):
    """Convert raw text to structured manual sections.
    Default: single LLM call with full raw_text + MANUALIZE_PROMPT.
    Exception: if len(raw_text) > HARD_LIMIT, uses window-based fallback internally."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")

    raw_text = doc["raw_text"]
    if not raw_text or raw_text.strip() == "" or raw_text.startswith("["):
        conn.close()
        error_msg = "Ï∂îÏ∂úÎêú ÌÖçÏä§Ìä∏Í∞Ä ÏóÜÍ±∞ÎÇò Ïú†Ìö®ÌïòÏßÄ ÏïäÏäµÎãàÎã§. Î®ºÏ†Ä 'Extract'Î•º ÏàòÌñâÌï¥ Ï£ºÏÑ∏Ïöî."
        if raw_text and raw_text.startswith("["):
            error_msg = f"ÌÖçÏä§Ìä∏ Ï∂îÏ∂úÏóê Î¨∏Ï†úÍ∞Ä ÏûàÏäµÎãàÎã§: {raw_text}"
        raise HTTPException(status_code=400, detail=error_msg)

    # Check if manual sections already exist (return cached if not forced)
    if not force:
        cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
        existing = cursor.fetchall()
        if existing:
            sections = {row["section_name"]: row["section_text"] for row in existing}
            conn.close()
            return {
                "success": True,
                "doc_id": doc_id,
                "sections": list(sections.keys()),
                "section_details": sections,
                "llm_used": False,
                "cached": True,
                "todo_questions": []
            }

    # LLM is required
    if not is_llm_available():
        conn.close()
        raise HTTPException(status_code=503, detail="LLMÏù¥ ÌôúÏÑ±ÌôîÎêòÏßÄ ÏïäÏïòÏäµÎãàÎã§. .env ÏÑ§Ï†ïÏùÑ ÌôïÏù∏ÌïòÏÑ∏Ïöî.")

    # Quick connectivity check before starting long operation
    llm_ok, llm_detail = ping_llm(timeout=5.0)
    if not llm_ok:
        conn.close()
        raise HTTPException(status_code=503, detail=llm_detail)

    raw_text_chars = len(raw_text)
    mode = get_llm_mode()
    window_count = len(_split_by_headings(raw_text))

    # ÏßÑÌñâÎ•† Ï¥àÍ∏∞Ìôî (sections: ÏôÑÎ£åÎêú ÏÑπÏÖòÏùÑ Ï¶âÏãú ÌîÑÎ°†Ìä∏Ïóê Ïä§Ìä∏Î¶¨Î∞ç)
    _manualize_progress[doc_id] = {"done": 0, "total": window_count, "sections": {}}

    try:
        if mode == "remote":
            # Remote: Ìï≠ÏÉÅ ÏÑπÏÖò Î∂ÑÌï† ‚Üí ÏàúÏ∞® Ï≤òÎ¶¨ (ÎÇ¥Ïö© Î≥¥Ï°¥)
            print(f"[MANUALIZE] Remote sequential: {raw_text_chars} chars, {window_count} sections for {doc_id}")
            sections_map = _manualize_with_window(raw_text, doc_id)
        elif raw_text_chars > MANUALIZE_HARD_LIMIT:
            # Local + ÎåÄÌòï Î¨∏ÏÑú: ÏúàÎèÑÏö∞ Î≥ëÎ†¨ Ï≤òÎ¶¨
            print(f"[MANUALIZE] Local parallel: {raw_text_chars} chars > {MANUALIZE_HARD_LIMIT}")
            sections_map = _manualize_with_window(raw_text, doc_id)
        else:
            # Local + ÏÜåÌòï Î¨∏ÏÑú: Îã®Ïùº Ìò∏Ï∂ú
            print(f"[MANUALIZE] Local single: {raw_text_chars} chars for {doc_id}")
            sections_map = _manualize_single(raw_text, doc_id)
            _manualize_progress[doc_id] = {"done": 1, "total": 1, "sections": dict(sections_map)}

    except HTTPException:
        raise
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=502, detail=f"Manualize Ïã§Ìå®: {str(e)}")

    # Check if cancelled
    if _manualize_progress.get(doc_id, {}).get("cancelled"):
        _manualize_progress.pop(doc_id, None)
        conn.close()
        raise HTTPException(status_code=499, detail="ÏÇ¨Ïö©ÏûêÍ∞Ä ManualizeÎ•º Ï∑®ÏÜåÌñàÏäµÎãàÎã§.")

    # Save sections (gate_status stays NULL ‚Äî user triggers Gate manually)
    cursor.execute("DELETE FROM manual_sections WHERE doc_id = %s", (doc_id,))
    cursor.execute("UPDATE documents SET completed_phases = NULL WHERE doc_id = %s", (doc_id,))
    sec_counter = 0
    for section_name, section_text in sections_map.items():
        sec_counter += 1
        if not section_name or not section_name.strip():
            section_name = f"## Í∏∞ÌÉÄ ({sec_counter})"
        section_id = f"sec_{uuid.uuid4().hex[:8]}"
        text_val = section_text if section_text else "Ï†ïÎ≥¥ ÏóÜÏùå"
        cursor.execute(
            """INSERT INTO manual_sections
               (section_id, doc_id, section_name, section_text, ai_text,
                gate_status, gate_score, gate_reasons_json, gate_stale)
               VALUES (%s, %s, %s, %s, %s, NULL, NULL, NULL, 0)""",
            (section_id, doc_id, section_name, text_val, text_val)
        )

    cursor.execute("UPDATE documents SET updated_at = %s WHERE doc_id = %s",
                   (datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()

    return {
        "success": True,
        "doc_id": doc_id,
        "sections": list(sections_map.keys()),
        "section_details": sections_map,
        "todo_questions": [],
        "change_summary": "",
        "pii_handling": {},
        "gate_results": {},
        "llm_used": True,
        "raw_text_chars": raw_text_chars,
        "window_count": window_count,
    }


def _manualize_single(raw_text: str, doc_id: str) -> dict:
    """Single LLM call with full raw_text. Returns {section_name: section_text}."""
    content = call_llm(MANUALIZE_PROMPT.format(raw_text=raw_text), temperature=0.3)
    if not content:
        raise Exception("LLM ÏùëÎãµÏù¥ ÎπÑÏñ¥ÏûàÏäµÎãàÎã§.")

    json_match = re.search(r'\{[\s\S]*\}', content)
    if not json_match:
        raise Exception("LLM ÏùëÎãµÏóêÏÑú JSONÏùÑ Ï∞æÏùÑ Ïàò ÏóÜÏäµÎãàÎã§.")

    parsed = _clean_llm_json(json_match.group())
    return _flatten_manualize_json(parsed)


def _split_by_headings(raw_text: str) -> list:
    """Split raw_text into sections by heading patterns.
    Returns list of (start_pos, section_text) tuples."""
    # Match common heading patterns: numbered (1. 2.1 Ï†ú3Ï°∞ etc.), markdown (#), or ALL-CAPS lines
    heading_pattern = re.compile(
        r'^(?:'
        r'#{1,4}\s+'                          # Markdown headings
        r'|Ï†ú?\s*\d+[Ï°∞Ìï≠Ïû•Ï†àÌé∏]\s*'           # Î≤ïÎ•†/Í∑úÏ†ï Ïä§ÌÉÄÏùº (Ï†ú1Ï°∞, Ï†ú2Ïû• Îì±)
        r'|\d+(?:\.\d+)*[\.\)]\s+'            # Numbered (1. 2.1. 3.2.1)
        r'|[Í∞Ä-Ìû£]{1,2}[\.\)]\s+'             # Korean bullets (Í∞Ä. ÎÇò.)
        r'|[IVX]+[\.\)]\s+'                   # Roman numerals
        r'|[A-Z][A-Z\s]{4,}$'                # ALL-CAPS lines (5+ chars)
        r')',
        re.MULTILINE
    )

    # Find all heading positions
    positions = [m.start() for m in heading_pattern.finditer(raw_text)]

    if not positions:
        # No headings found ‚Üí return whole text as one section
        return [(0, raw_text)]

    # Ensure we start from 0
    if positions[0] != 0:
        positions.insert(0, 0)

    # Build sections
    sections = []
    for i, pos in enumerate(positions):
        end = positions[i + 1] if i + 1 < len(positions) else len(raw_text)
        sections.append((pos, raw_text[pos:end]))

    return sections


def _group_sections_into_windows(sections: list, window_size: int) -> list:
    """Group heading-based sections into windows that fit within window_size.
    Each window is a string. If a single section exceeds window_size, it is sent alone."""
    windows = []
    current_window = ""

    for _pos, section_text in sections:
        # If adding this section would exceed limit
        if current_window and len(current_window) + len(section_text) > window_size:
            windows.append(current_window)
            current_window = section_text
        else:
            current_window += section_text

    if current_window:
        windows.append(current_window)

    return windows


def _process_one_window(window_text: str, idx: int, total: int) -> dict:
    """Process a single window. Returns {section_name: section_text} or empty dict."""
    print(f"[MANUALIZE] Processing window {idx + 1}/{total} ({len(window_text)} chars)...")
    try:
        content = call_llm(MANUALIZE_PROMPT.format(raw_text=window_text), temperature=0.3)
        if not content:
            return {}
        json_match = re.search(r'\{[\s\S]*\}', content)
        if not json_match:
            return {}
        parsed = _clean_llm_json(json_match.group())
        return _flatten_manualize_json(parsed)
    except Exception as e:
        print(f"[MANUALIZE] Window {idx + 1} failed: {e}")
        return {}


def _merge_sections(all_sections: dict, new_sections: dict, idx: int) -> None:
    """Merge new_sections into all_sections. Dedup by exact match, index on conflict."""
    for name, text in new_sections.items():
        if name in all_sections:
            if all_sections[name].strip() == text.strip():
                continue
            name = f"{name} ({idx + 1})"
        all_sections[name] = text


def _manualize_with_window(raw_text: str, doc_id: str) -> dict:
    """Large document processing. Mode-aware:
    - Local: group into windows ‚Üí parallel (concurrent threads)
    - Remote: heading sections ‚Üí sequential (ÏùëÎãµ Î∞õÍ≥† Îã§Ïùå Ï†ÑÏÜ°)"""
    sections = _split_by_headings(raw_text)
    print(f"[MANUALIZE] Found {len(sections)} heading sections for {doc_id}")

    mode = get_llm_mode()

    if mode == "remote":
        # Remote: ÏÑπÏÖòÎ≥Ñ ÏàúÏ∞® Ï≤òÎ¶¨ (ÏûëÏùÄ ÏûÖÎ†• ‚Üí ÎÇ¥Ïö© Î≥¥Ï°¥ Ïö∞Ïàò)
        total = len(sections)
        print(f"[MANUALIZE] Remote sequential: {total} sections for {doc_id}")
        _manualize_progress[doc_id] = {"done": 0, "total": total, "sections": {}}
        all_sections = {}
        for i, (_pos, section_text) in enumerate(sections):
            # Check cancellation before each section
            if _manualize_progress.get(doc_id, {}).get("cancelled"):
                print(f"[MANUALIZE] Cancelled at section {i}/{total} for {doc_id}")
                break
            if not section_text.strip():
                _manualize_progress[doc_id]["done"] = i + 1
                continue
            result = _process_one_window(section_text, i, total)
            _merge_sections(all_sections, result, i)
            _manualize_progress[doc_id]["done"] = i + 1
            _manualize_progress[doc_id]["sections"] = dict(all_sections)
    else:
        # Local: ÏúàÎèÑÏö∞ Î¨∂Ïñ¥ÏÑú Î≥ëÎ†¨ Ï≤òÎ¶¨ (ÏßÑÌñâÎ•†ÏùÄ ÏÑπÏÖò Ïàò Í∏∞Ï§ÄÏúºÎ°ú ÌëúÏãú)
        from concurrent.futures import ThreadPoolExecutor, as_completed
        windows = _group_sections_into_windows(sections, MANUALIZE_WINDOW_SIZE)
        total_windows = len(windows)
        total_sections = len(sections)
        print(f"[MANUALIZE] Local parallel: {total_windows} windows ({total_sections} sections) for {doc_id}")
        _manualize_progress[doc_id] = {"done": 0, "total": total_sections, "sections": {}}

        # Compute how many sections each window covers (mirrors _group_sections_into_windows logic)
        _sections_per_window = []
        _cur_len = 0
        _cur_cnt = 0
        for _pos, _st in sections:
            if _cur_len > 0 and _cur_len + len(_st) > MANUALIZE_WINDOW_SIZE:
                _sections_per_window.append(_cur_cnt)
                _cur_cnt = 1
                _cur_len = len(_st)
            else:
                _cur_cnt += 1
                _cur_len += len(_st)
        if _cur_cnt > 0:
            _sections_per_window.append(_cur_cnt)

        all_sections = {}
        if total_windows == 1:
            result = _process_one_window(windows[0], 0, 1)
            _merge_sections(all_sections, result, 0)
            _manualize_progress[doc_id]["done"] = total_sections
            _manualize_progress[doc_id]["sections"] = dict(all_sections)
        else:
            with ThreadPoolExecutor(max_workers=total_windows) as executor:
                futures = {
                    executor.submit(_process_one_window, w, i, total_windows): i
                    for i, w in enumerate(windows)
                }
                for future in as_completed(futures):
                    i = futures[future]
                    result = future.result()
                    _merge_sections(all_sections, result, i)
                    done_inc = _sections_per_window[i] if i < len(_sections_per_window) else 1
                    _manualize_progress[doc_id]["done"] = _manualize_progress[doc_id]["done"] + done_inc
                    _manualize_progress[doc_id]["sections"] = dict(all_sections)

    # ÏôÑÎ£å ÌõÑ Ï†ïÎ¶¨
    _manualize_progress.pop(doc_id, None)

    if not all_sections:
        raise Exception("Î™®Îì† Ï≤òÎ¶¨Í∞Ä Ïã§Ìå®ÌñàÏäµÎãàÎã§.")

    return all_sections


# ============ STEP 2a: SPLIT (progressive) ============

@router.post("/doc/{doc_id}/split")
def split_document(doc_id: str, force: bool = False):
    """Split document into chunks (rule-based). Returns chunk list for progressive manualize.
    If not force and chunks already exist, returns cached chunks."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")

    raw_text = doc["raw_text"]
    if not raw_text or raw_text.strip() == "" or raw_text.startswith("["):
        conn.close()
        raise HTTPException(status_code=400, detail="Ï∂îÏ∂úÎêú ÌÖçÏä§Ìä∏Í∞Ä ÏóÜÏäµÎãàÎã§. Î®ºÏ†Ä ExtractÎ•º ÏàòÌñâÌï¥ Ï£ºÏÑ∏Ïöî.")

    # Return cached chunks if not force
    if not force:
        cursor.execute("SELECT chunk_id, chunk_index, split_basis, notes, raw_chunk FROM doc_chunks WHERE doc_id = %s ORDER BY chunk_index", (doc_id,))
        existing = cursor.fetchall()
        if existing:
            chunks = []
            chunk_dicts = []
            for row in existing:
                preview = (row["raw_chunk"] or "")[:200]
                heading = row["notes"] or ""
                if not heading and row["raw_chunk"]:
                    first_line = row["raw_chunk"].split("\n", 1)[0].strip().lstrip("#").strip()
                    heading = first_line[:60] if first_line else f"Chunk {row['chunk_index'] + 1}"
                chunks.append({
                    "chunk_id": row["chunk_id"],
                    "chunk_index": row["chunk_index"],
                    "heading": heading,
                    "preview": preview,
                    "char_count": len(row["raw_chunk"] or ""),
                })
                chunk_dicts.append({
                    "chunk_id": row["chunk_id"],
                    "chunk_index": row["chunk_index"],
                    "notes": row["notes"] or "",
                    "raw_chunk": row["raw_chunk"] or "",
                })
            batches = _build_batches(chunk_dicts, doc_id)
            batch_info = [{"batch_id": b["batch_id"], "chunk_ids": b["chunk_ids"],
                           "oversize": [o["item_id"] for o in b.get("oversize_sections", [])]}
                          for b in batches]
            conn.close()
            return {"success": True, "doc_id": doc_id, "chunks": chunks, "batches": batch_info, "cached": True}

    # Split
    doc_chunks = _split_document(raw_text, doc_id, cursor)
    conn.commit()

    chunks = []
    for chunk in doc_chunks:
        raw_chunk = chunk.get("raw_chunk", "")
        heading = chunk.get("notes", "") or ""
        if not heading and raw_chunk:
            first_line = raw_chunk.split("\n", 1)[0].strip().lstrip("#").strip()
            heading = first_line[:60] if first_line else f"Chunk {chunk.get('chunk_index', 0) + 1}"
        chunks.append({
            "chunk_id": chunk["chunk_id"],
            "chunk_index": chunk.get("chunk_index", 0),
            "heading": heading,
            "preview": raw_chunk[:200],
            "char_count": len(raw_chunk),
        })

    batches = _build_batches(doc_chunks, doc_id)
    batch_info = [{"batch_id": b["batch_id"], "chunk_ids": b["chunk_ids"],
                   "oversize": [o["item_id"] for o in b.get("oversize_sections", [])]}
                  for b in batches]

    conn.close()
    return {"success": True, "doc_id": doc_id, "chunks": chunks, "batches": batch_info, "cached": False}


# ============ STEP 2b: MANUALIZE-CHUNK (progressive) ============

@router.post("/doc/{doc_id}/manualize-chunk/{chunk_id}")
def manualize_single_chunk(doc_id: str, chunk_id: str):
    """Manualize a single chunk. Returns section data immediately.
    Saves to manual_sections (appends, does not delete existing)."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT chunk_id, chunk_index, raw_chunk, notes FROM doc_chunks WHERE chunk_id = %s AND doc_id = %s", (chunk_id, doc_id))
    chunk_row = cursor.fetchone()
    if not chunk_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Chunk not found")

    if not is_llm_available():
        conn.close()
        raise HTTPException(status_code=503, detail="LLM ÎπÑÌôúÏÑ±")

    chunk = {
        "chunk_id": chunk_row["chunk_id"],
        "chunk_index": chunk_row["chunk_index"],
        "raw_chunk": chunk_row["raw_chunk"] or "",
        "notes": chunk_row["notes"] or "",
    }

    result = _manualize_chunk(chunk, doc_id)

    section_name = result["section_name"]
    if not section_name or not section_name.strip():
        section_name = f"## Í∏∞ÌÉÄ (chunk-{chunk_id[:8]})"
    section_text = result["section_text"] or "Ï†ïÎ≥¥ ÏóÜÏùå"
    evidence = result.get("evidence_spans", [])

    # Save section (check if already exists for this chunk)
    cursor.execute("SELECT section_id FROM manual_sections WHERE doc_id = %s AND source_chunk_id = %s", (doc_id, chunk_id))
    existing = cursor.fetchone()
    if existing:
        cursor.execute(
            """UPDATE manual_sections SET section_name = %s, section_text = %s,
               evidence_json = %s, updated_at = %s WHERE section_id = %s""",
            (section_name, section_text,
             json.dumps(evidence, ensure_ascii=False),
             datetime.now().isoformat(), existing["section_id"])
        )
        section_id = existing["section_id"]
    else:
        section_id = f"sec_{uuid.uuid4().hex[:8]}"
        cursor.execute(
            """INSERT INTO manual_sections
               (section_id, doc_id, section_name, section_text,
                source_chunk_id, evidence_json, merge_status)
               VALUES (%s, %s, %s, %s, %s, %s, 'BODY')""",
            (section_id, doc_id, section_name, section_text,
             chunk_id, json.dumps(evidence, ensure_ascii=False))
        )

    cursor.execute("UPDATE documents SET updated_at = %s WHERE doc_id = %s",
                   (datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()

    evidence_matched = sum(1 for s in evidence if s.get("char_start", -1) >= 0)

    return {
        "success": result.get("success", True),
        "chunk_id": chunk_id,
        "section_id": section_id,
        "section_name": section_name,
        "section_text": section_text,
        "evidence_spans": evidence,
        "evidence_matched": evidence_matched,
        "evidence_total": len(evidence),
    }


# ============ STEP 2c: MANUALIZE-BATCH (progressive) ============

@router.post("/doc/{doc_id}/manualize-batch/{batch_id}")
def manualize_batch_endpoint(doc_id: str, batch_id: str):
    """Manualize a batch of chunks in a single LLM call.
    Returns section data for all items in the batch."""
    conn = get_connection()
    cursor = conn.cursor()

    # Load chunks for this batch
    cursor.execute("SELECT chunk_id, chunk_index, notes, raw_chunk FROM doc_chunks WHERE doc_id = %s ORDER BY chunk_index", (doc_id,))
    all_chunks = [dict(row) for row in cursor.fetchall()]
    if not all_chunks:
        conn.close()
        raise HTTPException(status_code=404, detail="No chunks found")

    # Rebuild batches to find the requested one
    batches = _build_batches(all_chunks, doc_id)
    target_batch = None
    for b in batches:
        if b["batch_id"] == batch_id:
            target_batch = b
            break
    if not target_batch:
        conn.close()
        raise HTTPException(status_code=404, detail=f"Batch {batch_id} not found")

    if not is_llm_available():
        conn.close()
        raise HTTPException(status_code=503, detail="LLM ÎπÑÌôúÏÑ±")

    # Call LLM with batch prompt
    results = _manualize_batch(target_batch, doc_id)

    # Build chunk_id ‚Üí raw_chunk map for evidence indexing
    chunk_map = {c["chunk_id"]: c["raw_chunk"] for c in all_chunks}

    # Save each result to manual_sections
    saved = []
    for item in target_batch["items"]:
        chunk_id = item["item_id"]
        # Find matching result
        result = None
        for r in results:
            if r.get("item_id") == chunk_id:
                result = r
                break
        if not result:
            # Fallback: use first unmatched result or raw text
            if results:
                result = results.pop(0)
            else:
                result = {
                    "section_name": item["section_title"] or f"Chunk {chunk_id}",
                    "section_text": item["section_text"][:3000],
                    "evidence_spans": [],
                }

        section_name = result.get("section_name", item["section_title"] or f"Chunk {chunk_id}")
        if not section_name or not section_name.strip():
            section_name = f"## Í∏∞ÌÉÄ (chunk-{chunk_id[:8]})"
        section_text = result.get("section_text", "") or "Ï†ïÎ≥¥ ÏóÜÏùå"
        evidence_spans = result.get("evidence_spans", [])

        # Index evidence spans
        raw_chunk = chunk_map.get(chunk_id, "")
        indexed_spans = _index_evidence_spans(evidence_spans, raw_chunk)

        # Save/update
        cursor.execute("SELECT section_id FROM manual_sections WHERE doc_id = %s AND source_chunk_id = %s", (doc_id, chunk_id))
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """UPDATE manual_sections SET section_name = %s, section_text = %s,
                   evidence_json = %s, updated_at = %s WHERE section_id = %s""",
                (section_name, section_text,
                 json.dumps(indexed_spans, ensure_ascii=False),
                 datetime.now().isoformat(), existing["section_id"]))
            section_id = existing["section_id"]
        else:
            section_id = f"sec_{uuid.uuid4().hex[:8]}"
            cursor.execute(
                """INSERT INTO manual_sections
                   (section_id, doc_id, section_name, section_text,
                    source_chunk_id, evidence_json, merge_status)
                   VALUES (%s, %s, %s, %s, %s, %s, 'BODY')""",
                (section_id, doc_id, section_name, section_text,
                 chunk_id, json.dumps(indexed_spans, ensure_ascii=False)))

        evidence_matched = sum(1 for s in indexed_spans if s.get("char_start", -1) >= 0)
        saved.append({
            "chunk_id": chunk_id,
            "section_id": section_id,
            "section_name": section_name,
            "section_text": section_text,
            "evidence_spans": indexed_spans,
            "evidence_matched": evidence_matched,
            "evidence_total": len(indexed_spans),
        })

    # Handle oversize sections
    oversize_errors = []
    for o in target_batch.get("oversize_sections", []):
        oversize_errors.append({
            "chunk_id": o["item_id"],
            "code": "OVERSIZE_SECTION",
            "message": f"ÏÑπÏÖòÏù¥ {o['char_len']}ÏûêÎ°ú Î∞∞Ïπò ÌïúÎèÑ({MAX_BATCH_CHARS}Ïûê) Ï¥àÍ≥º",
        })

    cursor.execute("UPDATE documents SET updated_at = %s WHERE doc_id = %s",
                   (datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()

    return {
        "success": True,
        "batch_id": batch_id,
        "sections": saved,
        "errors": oversize_errors,
    }


# ============ STEP 2: SECTION GATE (per-section AI check) ============

GATE_CHECK_PROMPT = """ÎãπÏã†ÏùÄ RAG Î∞òÏòÅ Ï†Ñ, Îß§Îâ¥Ïñº ÏÑπÏÖò ÌÖçÏä§Ìä∏(section_text)Ïùò ÌíàÏßà/Î¶¨Ïä§ÌÅ¨Î•º ÌåêÏ†ïÌïòÎäî QA Í≤åÏù¥Ìä∏ÏûÖÎãàÎã§.
Ï§ëÏöî: ÎÇ¥Ïö©ÏùÑ ÏÉàÎ°ú ÏûëÏÑ±ÌïòÍ±∞ÎÇò Í≥†ÏπòÏßÄ ÎßêÍ≥†, Ïò§ÏßÅ 'Í≤ÄÏ¶ù Í≤∞Í≥º'Îßå Ï∂úÎ†•ÌïòÏÑ∏Ïöî.

[ÏûÖÎ†•]
section_text: {section_text}
raw_text: {raw_text}

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[ÌåêÏ†ï ÏÉÅÌÉú]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
PASS: Î∞îÎ°ú RAG Î∞òÏòÅ Í∞ÄÎä•
NEED_FIX: ÏÇ¨ÎûåÏù¥ ÏàòÏ†ï/Î≥¥Í∞ï ÌõÑ Î∞òÏòÅ Í∂åÏû•
BLOCK: RAG Î∞òÏòÅ Í∏àÏßÄ(Î≥¥Ïïà/Ïã¨Í∞Å Ï∂©Îèå/ÎåÄÎüâ Ïú†Ïã§)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[üö® ÏµúÏö∞ÏÑ† ÏõêÏπô]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
- Ï∂îÏ∏° Í∏àÏßÄ.
- Î™®Îì† ÏßÄÏ†ÅÏùÄ section_text ÎòêÎäî raw_textÏùò **Ï¶ùÍ±∞ Î¨∏ÏûêÏó¥ Í∏∞Î∞ò**Ïù¥Ïñ¥Ïïº Ìï®.
- location_hintÏóê Ï¶ùÍ±∞Î•º Ï†úÏãúÌï† Ïàò ÏóÜÏúºÎ©¥ reasonsÏóê ÎÑ£ÏßÄ ÎßàÏÑ∏Ïöî.
- "Í∞ÄÎä•ÏÑ±", "Ï∂îÏ†ï", "ÏïÑÎßà" Í∞ôÏùÄ ÌëúÌòÑ Í∏∞Î∞ò ÏßÄÏ†Å Í∏àÏßÄ.
- Í∞ôÏùÄ ÎÇ¥Ïö©ÏùÑ Î∞òÎ≥µ/Í≥ºÏûâ ÏßÄÏ†ÅÌïòÏßÄ ÎßàÏÑ∏Ïöî(ÌïµÏã¨Îßå).

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Ïö¥ÏòÅ ÏïàÏ†ÑÌïÄ]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
- HIGH severityÎäî "Î™ÖÎ∞±ÌïòÍ≥† ÏßÅÏ†ëÏ†ÅÏù∏ Ï¶ùÍ±∞"Í∞Ä ÏûàÏùÑ ÎïåÎßå Î∂ÄÏó¨ÌïòÏÑ∏Ïöî.
  Ïï†Îß§ÌïòÍ±∞ÎÇò Ìï¥ÏÑù Ïó¨ÏßÄÍ∞Ä ÏûàÏúºÎ©¥ HIGHÎ°ú Ïò¨Î¶¨ÏßÄ ÎßêÍ≥† MEDIUMÏúºÎ°ú ÎëêÏÑ∏Ïöî(Í≥ºÎèÑÌïú BLOCK Î∞©ÏßÄ).
- HIGHÎäî ÏùºÎ∞ò Ìï¥ÏÑùÏù¥ ÏïÑÎãàÎùº, ÏÇ¨ÎûåÏù¥ Î≥¥ÏïÑÎèÑ Î™ÖÎ∞±Ìïú ÏúÑÎ∞ò/Ï∂©ÎèåÎ°ú Ïù∏ÏãùÎêòÎäî Í≤ΩÏö∞ÏóêÎßå Î∂ÄÏó¨ÌïòÏÑ∏Ïöî.
- scoreÎäî Ï†ïÍµêÌïú ÌèâÍ∞ÄÍ∞Ä ÏïÑÎãàÎùº status Íµ¨Í∞Ñ(PASS/NEED_FIX/BLOCK)ÏùÑ Î∞òÏòÅÌïòÎäî ÎåÄÎûµÏ†Å Í∞íÏûÖÎãàÎã§.
  Ïö¥ÏòÅ ÌåêÎã®ÏùÄ scoreÍ∞Ä ÏïÑÎãàÎùº statusÎ•º Ïö∞ÏÑ†ÌïòÏÑ∏Ïöî.

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Í≤ÄÏ¶ù Ìï≠Î™©]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

1) PII_RISK (BLOCK Ïö∞ÏÑ† ‚Äî Ïã§Ï†ú Í∞íÏùº ÎïåÎßå)
Îã§Ïùå **Ïã§Ï†ú Í∞í**Ïù¥ ÎßàÏä§ÌÇπ ÏóÜÏù¥ ÏûàÏùÑ ÎïåÎßå Î¨∏Ï†ú:
- Ï£ºÎØºÎ≤àÌò∏ Ìå®ÌÑ¥
- Í≥ÑÏ¢å/IDÏÑ± Í∏¥ Ïà´ÏûêÏó¥(10~16ÏûêÎ¶¨)

PII_RISKÍ∞Ä ÏïÑÎãå Í≤É (ÏòàÏô∏):
- Í≥†Í∞ùÏÑºÌÑ∞/ÎåÄÌëú Ï†ÑÌôîÎ≤àÌò∏ (02-xxxx-xxxx, 1588-xxxx Îì±): Í≥µÍ∞ú Ïó∞ÎùΩÏ≤òÏù¥ÎØÄÎ°ú PII ÏïÑÎãò
- ÏÑúÎπÑÏä§ ÏïàÎÇ¥Ïö© Ïù¥Î©îÏùº/URL: Í≥µÍ∞ú Ï†ïÎ≥¥Ïù¥ÎØÄÎ°ú PII ÏïÑÎãò
- Í∞úÏù∏ Ìú¥ÎåÄÌè∞(010-xxxx-xxxx)ÎèÑ Îß§Îâ¥ÏñºÏóê ÏùòÎèÑÏ†ÅÏúºÎ°ú Í∏∞Ïû¨Îêú Í≤ÉÏù¥Î©¥ PII ÏïÑÎãò
- "Ï†ÑÌôîÎ≤àÌò∏/Ïù¥Î©îÏùº/Ï£ºÏÜå" Í∞ôÏùÄ ÌïÑÎìúÎ™ÖÏùÄ PII ÏïÑÎãò

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

2) CONFLICT
ÎèôÏùº Ï°∞Í±¥/ÎåÄÏÉÅ/ÏÉÅÌô©ÏóêÏÑú **ÏßÅÏ†ë Ï∂©Îèå**ÌïòÎäî Í∑úÏπôÎßå Ìï¥Îãπ.
(Í∞ÄÎä• vs Î∂àÍ∞Ä, ÌóàÏö© vs Í∏àÏßÄ Îì±)
location_hint: Ï∂©ÎèåÌïòÎäî Î¨∏Ïû• Ï§ë ÌïòÎÇòÎ•º section_textÏóêÏÑú **ÏõêÎ¨∏ Í∑∏ÎåÄÎ°ú** Î≥µÏÇ¨ (10~60Ïûê)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

3) MISSING / AMBIGUOUS
Ïã§Ìñâ Í∑úÏπôÏù∏Îç∞ ÌïÑÏàòÍ∞í(Í∏∞Ìïú/Í∏àÏï°/Ï°∞Í±¥/Ï±ÑÎÑê/Îã¥Îãπ) ÏóÜÏúºÎ©¥ MISSING.
Î™®Ìò∏ ÌëúÌòÑ Î∞òÎ≥µ Ïãú AMBIGUOUS:
(Ï†ÅÎãπÌûà/Í∞ÄÎä•ÌïòÎ©¥/ÏÉÅÌô©Ïóê Îî∞Îùº/ÌòëÏùò ÌõÑ/ÌïÑÏöîÏãú)
ÏÜåÍ∞ú/Í∞úÏöî ÏÑπÏÖòÏóêÎäî Í≥ºÎèÑ Ï†ÅÏö© Í∏àÏßÄ.

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

4) HALLUCINATION_RISK
section_textÏóê ÏàòÏπò/Í∏∞Í∞Ñ/Í∏àÏï°/ÏùòÎ¨¥ Í∑úÏπôÏù¥ ÏûàÎäîÎç∞,
raw_textÏóêÏÑú **ÎèôÏùº Ï£ºÏ†ú Í¥ÄÎ†® ÌëúÌòÑ Ï†ÑÎ∞ò**ÏùÑ Ï∞æÍ∏∞ Ïñ¥Î†§Ïö∏ ÎïåÎßå Ìï¥Îãπ.
- Îã®Ïùº ÌÇ§ÏõåÎìú Î∂ÄÏû¨Î°ú ÌåêÎã® Í∏àÏßÄ.
- section_textÏóêÏÑú ÌïµÏã¨ ÌÇ§ÏõåÎìú 1~3Í∞úÎ•º Î®ºÏ†Ä ÎΩëÏïÑ ÎπÑÍµê.

ÏòàÏô∏ (HALLUCINATION_RISKÎ°ú ÌåêÏ†ïÌïòÏßÄ Îßê Í≤É):
- Ïó∞ÎùΩÏ≤ò(Ï†ÑÌôîÎ≤àÌò∏), Ïù¥Î©îÏùº, URL, ÏõπÏÇ¨Ïù¥Ìä∏ Ï£ºÏÜå: ÏÇ¨Ïö©ÏûêÍ∞Ä ÏùòÎèÑÏ†ÅÏúºÎ°ú Ï∂îÍ∞ÄÌñàÏùÑ Ïàò ÏûàÏùå
- "ÏßÄÏõê", "Î¨∏Ïùò", "ÏïàÎÇ¥" Îì± ÏÑúÎπÑÏä§ ÏïàÎÇ¥ ÏÑπÏÖòÏùò Ïó∞ÎùΩ Ï†ïÎ≥¥Îäî Ïã†Î¢∞
- [ÏÇ¨Ïö©Ïûê ÏßÅÏ†ë ÏàòÏ†ï Íµ¨Í∞Ñ]ÏúºÎ°ú ÌëúÏãúÎêú ÎÇ¥Ïö©

location_hint:
- section_textÏóêÏÑú Î¨∏Ï†ú ÎêòÎäî Î¨∏Ïû•/Î∂àÎ¶øÏùÑ **ÏõêÎ¨∏ Í∑∏ÎåÄÎ°ú** Î≥µÏÇ¨ (10~60Ïûê)

severity:
- Í∏∞Î≥∏ MEDIUM
- HIGHÎäî ÏÑúÎ°ú Îã§Î•∏ Íµ¨Ï≤¥ ÏàòÏπò/Í∏∞Í∞Ñ/Í∏àÏï°/ÏùòÎ¨¥ Í∑úÏ†ïÏù¥ 2Í∞ú Ïù¥ÏÉÅ Ï∂îÍ∞ÄÎêú Í≤ΩÏö∞Îßå
  (Í∑∏Î¶¨Í≥† Í∑∏ Ï∂îÍ∞ÄÍ∞Ä Î™ÖÎ∞±Ìûà ÌôïÏù∏Îê† ÎïåÎßå)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

5) FORMAT
FORMATÏùÄ ÏïÑÎûòÏùº ÎïåÎßå:
- section_textÍ∞Ä "##"Î°ú ÏãúÏûëÌïòÏßÄ ÏïäÏùå
- ### Ìï≠Î™©Ïù¥ 0Í∞ú AND '- ' bulletÎèÑ 3Í∞ú ÎØ∏Îßå

ÏòàÏô∏:
### 0Í∞úÎùºÎèÑ '- ' bullet ‚â• 3Ïù¥Î©¥ FORMAT ÏïÑÎãò.

FORMAT Îã®ÎèÖ BLOCK Í∏àÏßÄ.
(ÌòïÏãù Î¨∏Ï†úÎäî ÏµúÎåÄ NEED_FIXÍπåÏßÄÎßå)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ

6) OMISSION (Ï†ïÎ≥¥ Ïú†Ïã§) ‚Äî Î∞©Ìñ•: raw_text ‚Üí section_textÎßå Ìï¥Îãπ
raw_textÏùò **Ïã§Ìñâ Ï†ïÎ≥¥**Í∞Ä section_textÏóê ÏóÜÏùÑ ÎïåÎßå OMISSION.
Ï†ÅÏö© ÎåÄÏÉÅ:
- Ï†àÏ∞® / Ï°∞Í±¥ / ÏòàÏô∏ / Í∏àÏï° / Í∏∞Ìïú / Îã¥Îãπ / Ï±ÑÎÑê
Ï†ÅÏö© Ï†úÏô∏:
- Î∞∞Í≤Ω ÏÑ§Î™Ö / ÌôçÎ≥¥ Î¨∏Íµ¨ / ÏòàÏãú/Î∂ÄÏó∞

‚ö†Ô∏è Î∞©Ìñ• Ï£ºÏùò:
- OMISSION = raw_textÏóê ÏûàÎäîÎç∞ section_textÏóê ÏóÜÎäî Í≤É
- section_textÏóê ÏûàÎäîÎç∞ raw_textÏóê ÏóÜÎäî Í≤ÉÏùÄ OMISSIONÏù¥ ÏïÑÎãò (‚Üí HALLUCINATION_RISK Í≤ÄÌÜ† ÎåÄÏÉÅ)
- section_textÏóêÎßå Ï°¥Ïû¨ÌïòÎäî Ïó∞ÎùΩÏ≤ò/URL/Ïù¥Î©îÏùº Îì±ÏùÑ OMISSIONÏúºÎ°ú Î∂ÑÎ•òÌïòÏßÄ ÎßàÏÑ∏Ïöî.
- [ÏÇ¨Ïö©ÏûêÍ∞Ä ÏßÅÏ†ë ÏÇ≠Ï†úÌïú ÎÇ¥Ïö©]ÏúºÎ°ú ÌëúÏãúÎêú ÎÇ¥Ïö©ÏùÄ OMISSIONÏúºÎ°ú ÌåêÏ†ïÌïòÏßÄ ÎßàÏÑ∏Ïöî.

ÌåêÏ†ï(Ï¶ùÍ±∞ ÌïÑÏàò):
- raw_textÏóêÏÑú ÎàÑÎùΩ ÏòàÏãú 2Í∞ú(ÏßßÏùÄ Î¨∏Ïû•/Î∂àÎ¶ø) Ï†úÏãú
- section_textÏóêÏÑú Ìï¥Îãπ ÎÇ¥Ïö©Ïù¥ ÏóÜÏùåÏùÑ Î™ÖÏãú

location_hint: ÎàÑÎùΩÎêú ÎÇ¥Ïö©Ïù¥ Ï∂îÍ∞ÄÎêòÏñ¥Ïïº Ìï† section_textÏùò **Ïù∏Ï†ë Ï§Ñ**ÏùÑ ÏõêÎ¨∏ Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨ (10~60Ïûê).
(Ïòà: raw_textÏóê 'ÎùΩÏª§ Ï¢ÖÎ•ò: ÏÜåÌòï 40x40cm'Í∞Ä ÎàÑÎùΩ ‚Üí section_textÏùò 'ÎùΩÏª§(ÏÇ¨Î¨ºÌï®)ÏÑ§Ï†ï' Ï§ÑÏùÑ Î≥µÏÇ¨)

ÎåÄÎüâ ÎàÑÎùΩ(ÌïµÏã¨ Ïã§ÌñâÏ†ïÎ≥¥ Îã§Ïàò ÎàÑÎùΩ) ÏãúÎßå BLOCK.
- "ÎåÄÎüâ ÎàÑÎùΩ"ÏùÄ Ïã§Ìñâ Ï†ïÎ≥¥ Ïú†Ìòï(Ï†àÏ∞®/Ï°∞Í±¥/ÏòàÏô∏/Í∏àÏï°/Í∏∞Ìïú/Îã¥Îãπ/Ï±ÑÎÑê) Ï§ë **2Í∞ÄÏßÄ Ïù¥ÏÉÅ Ïú†ÌòïÏóê ÎåÄÌï¥**
  raw_textÏóê Î™ÖÏãúÎêú Ïã§Ìñâ Ï†ïÎ≥¥Í∞Ä section_textÏóêÏÑú **Î∞òÎ≥µÏ†ÅÏúºÎ°ú ÎàÑÎùΩ**ÎêòÎäî Í≤ΩÏö∞Î•º ÏùòÎØ∏Ìï©ÎãàÎã§.
  (Îã®, Îã®Ïùº Ìï≠Î™© ÏàòÏ§ÄÏùò Í≤ΩÎØ∏ ÎàÑÎùΩÏùÄ HIGHÎ°ú Ïò¨Î¶¨ÏßÄ ÎßêÍ≥† MEDIUMÏúºÎ°ú ÎëêÏÑ∏Ïöî.)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Í≤ÄÏ¶ù Î∂àÏ∂©Î∂Ñ]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
PASS Í∏àÏßÄ ‚Üí NEED_FIX (INSUFFICIENT_EVIDENCE)

Ìï¥Îãπ Ï°∞Í±¥(Î∞òÎìúÏãú Í∑ºÍ±∞Î°ú ÏÑ§Î™Ö):
- raw_text < 800Ïûê AND (OMISSION/CONFLICT/HALLUCINATION_RISK) ÎåÄÏ°∞Ïóê ÌïÑÏöîÌïú Í∑ºÍ±∞Í∞Ä Î∂ÄÏ°±Ìï®
- section_text ÌïµÏã¨ ÌÇ§ÏõåÎìú(1~3Í∞ú)Î•º raw_textÏóêÏÑú Ï∞æÍ∏∞ Ïñ¥Î†§Ïõå Í∑ºÍ±∞ ÎåÄÏ°∞Í∞Ä ÏÇ¨Ïã§ÏÉÅ Î∂àÍ∞ÄÎä•Ìï®
- Í∑ºÍ±∞ Î∂ÄÏ°±ÏúºÎ°ú ÌåêÎã® Î∂àÍ∞Ä

location_hint: section_textÏóêÏÑú Í∑ºÍ±∞ Î∂ÄÏ°±Ìïú Î¨∏Ïû•ÏùÑ **ÏõêÎ¨∏ Í∑∏ÎåÄÎ°ú** Î≥µÏÇ¨ (10~60Ïûê)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Ï∂úÎ†• ÌòïÏãù]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
RFC8259 JSONÎßå Ï∂úÎ†•.

Ï†úÏïΩ:
- reasonsÎäî ÏµúÎåÄ 6Í∞úÍπåÏßÄÎßå Ï∂úÎ†•
- Í∞ôÏùÄ typeÏùÄ ÏµúÎåÄ 2Í∞úÍπåÏßÄÎßå Ï∂úÎ†•

{{
  "status": "PASS|NEED_FIX|BLOCK",
  "score": 0,
  "reasons": [
    {{
      "type": "PII_RISK|CONFLICT|MISSING|AMBIGUOUS|HALLUCINATION_RISK|FORMAT|OMISSION|INSUFFICIENT_EVIDENCE",
      "severity": "LOW|MEDIUM|HIGH",
      "message": "Î¨∏Ï†ú ÏÑ§Î™Ö",
      "location_hint": "section_textÏóêÏÑú Î¨∏Ï†úÍ∞Ä ÎêòÎäî Î∂ÄÎ∂ÑÏùÑ **Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨**ÌïòÏÑ∏Ïöî (10~60Ïûê). ÏõêÎ¨∏Í≥º Ìïú Í∏ÄÏûêÎùºÎèÑ Îã§Î•¥Î©¥ ÌïòÏù¥ÎùºÏù¥Ìä∏Í∞Ä Ïïà Îê©ÎãàÎã§. ÏÑ§Î™Ö/ÏÇ¨Ïú†Îäî Ïó¨Í∏∞Ïóê Ïì∞ÏßÄ ÎßàÏÑ∏Ïöî.",
      "fix_suggestion": "ÏàòÏ†ï Î∞©Î≤ï"
    }}
  ],
  "required_actions": []
}}

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Score Í∏∞Ï§Ä]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
PASS: 85~100
NEED_FIX: 50~84
BLOCK: 0~49

Í∞êÏ†ê Í∏∞Ï§Ä(100ÏóêÏÑú Í∞êÏ†ê):
HIGH = -40
MEDIUM = -20
LOW = -10

- scoreÎäî 0~100 Ï†ïÏàòÎ°ú ÌÅ¥Îû®ÌîÑ(Î≤îÏúÑÎ•º Î≤óÏñ¥ÎÇòÎ©¥ 0 ÎòêÎäî 100ÏúºÎ°ú Î≥¥Ï†ï)

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Status Í≤∞Ï†ï Í∑úÏπô]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
ÏïÑÎûò Ï°∞Í±¥Ïù¥ reasonsÏóê Ï°¥Ïû¨ÌïòÎ©¥ statusÎäî Î∞òÎìúÏãú Í∑∏Î†áÍ≤å Í≤∞Ï†ïÌï¥Ïïº Ìï©ÎãàÎã§(Í∞ïÏ†ú).

- PII_RISK(HIGH) ‚Üí BLOCK (Î¨¥Ï°∞Í±¥)
- CONFLICT(HIGH) ‚Üí BLOCK
- OMISSION(HIGH) ‚Üí BLOCK
- HALLUCINATION_RISK(HIGH) ‚Üí BLOCK

Ï∂îÍ∞Ä Í∑úÏπô:
- HIGH 1Í∞ú(ÏúÑ 4Í∞ú Ïú†Ìòï Ïô∏ Ìè¨Ìï®) ‚Üí ÏµúÏÜå NEED_FIX
- MEDIUM 2Í∞ú Ïù¥ÏÉÅ ‚Üí NEED_FIX
- LOWÎßå Ï°¥Ïû¨ ‚Üí PASS Í∞ÄÎä•
- FORMATÏùÄ BLOCK Í∏àÏßÄ Ïú†ÏßÄ

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[PASS ÏµúÏ¢Ö Ï°∞Í±¥]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
reasonsÍ∞Ä ÎπÑÏñ¥ ÏûàÏñ¥ÎèÑ ÏïÑÎûò Î™®Îëê ÎßåÏ°±:
- PII ÏóÜÏùå
- FORMAT ÌÜµÍ≥º
- Í≤ÄÏ¶ù Ï∂©Î∂Ñ(INSUFFICIENT_EVIDENCE ÏïÑÎãò)
- Î™ÖÎ∞±Ìïú ÎàÑÎùΩ/ÌôòÍ∞Å ÏóÜÏùå
"""


def _gate_section_internal(doc_id: str, section_name: str, cursor, raw_text_fallback: str = None):
    """Core gate logic reusable by gate_section() and gate-all.
    cursor must be from an open connection. Does NOT commit/close.
    raw_text_fallback: optional pre-fetched raw_text to avoid redundant query."""
    import time as _time
    t_start = _time.time()

    cursor.execute("SELECT section_id, section_text, source_chunk_id, ai_text, gate_reasons_json FROM manual_sections WHERE doc_id = %s AND section_name = %s",
                   (doc_id, section_name))
    sec = cursor.fetchone()
    if not sec:
        return {"success": False, "section_name": section_name, "error": "Section not found"}

    # Load previously dismissed reasons
    prev_dismissed = []
    try:
        prev_reasons = json.loads(sec["gate_reasons_json"] or "[]")
        prev_dismissed = [r for r in prev_reasons if r.get("dismissed")]
    except Exception:
        pass

    # Determine raw reference: prefer raw_chunk from doc_chunks if available
    raw_reference = ""
    if sec["source_chunk_id"]:
        cursor.execute("SELECT raw_chunk FROM doc_chunks WHERE chunk_id = %s", (sec["source_chunk_id"],))
        chunk_row = cursor.fetchone()
        if chunk_row and chunk_row["raw_chunk"]:
            raw_reference = chunk_row["raw_chunk"][:4000]

    if not raw_reference:
        if raw_text_fallback is not None:
            raw_reference = raw_text_fallback[:4000]
        else:
            cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
            doc = cursor.fetchone()
            raw_reference = (doc["raw_text"] or "")[:4000] if doc else ""

    # Detect user edits: diff section_text vs ai_text
    user_edits_note = ""
    ai_text = sec["ai_text"] or ""
    section_text = sec["section_text"] or ""
    if ai_text and section_text != ai_text:
        ai_lines = set(ai_text.strip().splitlines())
        sec_lines = set(section_text.strip().splitlines())
        sec_lines_list = section_text.strip().splitlines()
        # ÏÇ¨Ïö©ÏûêÍ∞Ä Ï∂îÍ∞ÄÌïú ÎùºÏù∏ (ai_textÏóê ÏóÜÍ≥† section_textÏóê ÏûàÏùå)
        added = [l for l in sec_lines_list if l.strip() and l not in ai_lines]
        # ÏÇ¨Ïö©ÏûêÍ∞Ä ÏÇ≠Ï†úÌïú ÎùºÏù∏ (ai_textÏóê ÏûàÍ≥† section_textÏóê ÏóÜÏùå)
        removed = [l for l in ai_text.strip().splitlines() if l.strip() and l not in sec_lines]
        parts = []
        if added:
            added_preview = "\n".join(added[:10])
            parts.append(f"[ÏÇ¨Ïö©ÏûêÍ∞Ä ÏßÅÏ†ë Ï∂îÍ∞ÄÌïú ÎÇ¥Ïö© ‚Äî HALLUCINATION_RISK Ï†úÏô∏ ÎåÄÏÉÅ]\n{added_preview}\n(ÏúÑ ÎÇ¥Ïö©ÏùÄ ÏÇ¨Ïö©ÏûêÍ∞Ä ÏùòÎèÑÏ†ÅÏúºÎ°ú Ï∂îÍ∞ÄÌïú Í≤ÉÏù¥ÎØÄÎ°ú raw_textÏóê ÏóÜÎçîÎùºÎèÑ HALLUCINATION_RISKÎ°ú ÌåêÏ†ïÌïòÏßÄ ÎßàÏÑ∏Ïöî.)")
        if removed:
            removed_preview = "\n".join(removed[:10])
            parts.append(f"[ÏÇ¨Ïö©ÏûêÍ∞Ä ÏßÅÏ†ë ÏÇ≠Ï†úÌïú ÎÇ¥Ïö© ‚Äî OMISSION Ï†úÏô∏ ÎåÄÏÉÅ]\n{removed_preview}\n(ÏúÑ ÎÇ¥Ïö©ÏùÄ ÏÇ¨Ïö©ÏûêÍ∞Ä ÏùòÎèÑÏ†ÅÏúºÎ°ú ÏÇ≠Ï†úÌïú Í≤ÉÏù¥ÎØÄÎ°ú section_textÏóê ÏóÜÎçîÎùºÎèÑ OMISSIONÏúºÎ°ú ÌåêÏ†ïÌïòÏßÄ ÎßàÏÑ∏Ïöî.)")
        if parts:
            user_edits_note = "\n\n" + "\n\n".join(parts)

    # Build dismissed note for LLM prompt
    dismissed_note = ""
    if prev_dismissed:
        dismissed_items = []
        for d in prev_dismissed:
            dismissed_items.append(f"- [{d.get('type','?')}] {d.get('description','')[:100]}")
        dismissed_note = "\n\n[Ïù¥Ï†Ñ Í≤ÄÏ¶ùÏóêÏÑú ÏÇ¨Ïö©ÏûêÍ∞Ä Î¨¥Ïãú(dismiss)Ìïú Ïù¥Ïäà - ÎèôÏùº Ïù¥ÏäàÎ•º Îã§Ïãú Î≥¥Í≥†ÌïòÏßÄ ÎßàÏÑ∏Ïöî]\n" + "\n".join(dismissed_items)

    gate_result = {"status": "PASS", "score": 100, "reasons": [], "required_actions": []}

    if is_llm_available():
        try:
            prompt_text = GATE_CHECK_PROMPT.format(
                section_text=section_text[:3000],
                raw_text=raw_reference
            ) + user_edits_note + dismissed_note
            print(f"[GATE_SECTION] calling LLM for '{section_name}' (user_edits_note: {len(user_edits_note)} chars, dismissed: {len(prev_dismissed)})")
            content = call_llm(prompt_text, temperature=0.3)
            print(f"[GATE_SECTION] LLM response for '{section_name}': {content[:300] if content else 'EMPTY'}")
            if content:
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    gate_result = _clean_llm_json(json_match.group())
                    print(f"[GATE_SECTION] result for '{section_name}': status={gate_result.get('status')}, reasons={len(gate_result.get('reasons', []))}")
        except Exception as e:
            print(f"[GATE_SECTION] LLM error: {e}")
    else:
        print(f"[GATE_SECTION] LLM not available, returning default PASS")

    # Post-process: restore dismissed flags for matching reasons
    new_reasons = gate_result.get("reasons", [])
    if prev_dismissed:
        dismissed_keys = {(d.get("type", ""), d.get("description", "")[:80]) for d in prev_dismissed}
        for r in new_reasons:
            key = (r.get("type", ""), r.get("description", "")[:80])
            if key in dismissed_keys:
                r["dismissed"] = True

    # Append previously dismissed reasons that LLM didn't re-report (keep them visible as dismissed)
    existing_keys = {(r.get("type", ""), r.get("description", "")[:80]) for r in new_reasons}
    for d in prev_dismissed:
        key = (d.get("type", ""), d.get("description", "")[:80])
        if key not in existing_keys:
            d["dismissed"] = True
            new_reasons.append(d)

    gate_result["reasons"] = new_reasons

    # Recalculate status from active (non-dismissed) reasons only
    active = [r for r in new_reasons if not r.get("dismissed")]
    if not active:
        gate_result["status"] = "PASS"
        gate_result["score"] = max(gate_result.get("score", 100), 85)
    else:
        has_high = any(r.get("severity") == "HIGH" for r in active)
        high_types = {r.get("type") for r in active if r.get("severity") == "HIGH"}
        block_types = {"PII_RISK", "CONFLICT", "OMISSION", "HALLUCINATION_RISK"}
        if has_high and high_types & block_types:
            gate_result["status"] = "BLOCK"
        elif has_high or sum(1 for r in active if r.get("severity") == "MEDIUM") >= 2:
            gate_result["status"] = "NEED_FIX"
        else:
            gate_result["status"] = "PASS"

    # Save gate result to manual_sections
    cursor.execute(
        "UPDATE manual_sections SET gate_status = %s, gate_score = %s, gate_reasons_json = %s, gate_stale = 0, updated_at = %s WHERE section_id = %s",
        (gate_result["status"], gate_result.get("score", 100),
         json.dumps(gate_result["reasons"], ensure_ascii=False),
         datetime.now().isoformat(), sec["section_id"])
    )

    gate_result["time_s"] = round(_time.time() - t_start, 2)
    return {"success": True, "section_name": section_name, **gate_result}


@router.post("/doc/{doc_id}/section/{section_name}/gate")
def gate_section(doc_id: str, section_name: str):
    """Run AI gate check on a single section.
    If source_chunk_id exists, use raw_chunk instead of full raw_text (chunk-based)."""
    conn = get_connection()
    cursor = conn.cursor()
    result = _gate_section_internal(doc_id, section_name, cursor)
    if not result.get("success"):
        conn.close()
        raise HTTPException(status_code=404, detail=result.get("error", "Section not found"))
    conn.commit()
    conn.close()
    return result


@router.post("/doc/{doc_id}/gate-all")
def gate_all(doc_id: str):
    """Run AI gate check on ALL sections of a document."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT section_name FROM manual_sections WHERE doc_id = %s", (doc_id,))
    sections = cursor.fetchall()
    if not sections:
        conn.close()
        raise HTTPException(status_code=400, detail="No sections found")

    # Pre-fetch raw_text once for fallback
    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    raw_text_fallback = (doc["raw_text"] or "") if doc else ""

    results = []
    pass_count = need_fix_count = block_count = 0
    for sec in sections:
        r = _gate_section_internal(doc_id, sec["section_name"], cursor, raw_text_fallback=raw_text_fallback)
        results.append(r)
        status = r.get("status", "PASS")
        if status == "PASS":
            pass_count += 1
        elif status == "NEED_FIX":
            need_fix_count += 1
        elif status == "BLOCK":
            block_count += 1

    conn.commit()
    conn.close()
    return {
        "success": True, "pass_count": pass_count,
        "need_fix_count": need_fix_count, "block_count": block_count,
        "results": results
    }


@router.post("/doc/{doc_id}/fill-all")
def fill_all(doc_id: str):
    """Îß•ÎùΩ Î≥¥Í∞ï(Fill)Îßå Ï†Ñ ÏÑπÏÖòÏóê Ïã§Ìñâ."""
    if not is_llm_available():
        raise HTTPException(status_code=503, detail="LLM ÏÇ¨Ïö© Î∂àÍ∞Ä")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")
    raw_text = doc["raw_text"] or ""
    raw_text_safe = raw_text[:4000] if raw_text else "(ÏõêÎ≥∏ Î¨∏ÏÑúÎ•º Ï∞æÏùÑ Ïàò ÏóÜÏäµÎãàÎã§. Í∏∞Ï°¥ ÌÖçÏä§Ìä∏Îßå Ï∞∏Í≥†ÌïòÏÑ∏Ïöî.)"

    cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    all_sections = cursor.fetchall()
    if not all_sections:
        conn.close()
        raise HTTPException(status_code=400, detail="No sections found")

    total = len(all_sections)
    done_count = 0
    failed_sections = []

    for sec in all_sections:
        section_name = sec["section_name"]
        try:
            qa_policy_text = "Q&AÎäî ÏÉàÎ°ú Ï∂îÍ∞ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî. Í∏∞Ï°¥ Q&AÎßå Ïú†ÏßÄ/Ï†ïÎ¶¨ÌïòÏÑ∏Ïöî."
            prompt = FILL_SECTION_TEXT_PROMPT_V3.format(
                section_text=sec["section_text"],
                raw_text=raw_text_safe,
                qa_policy_text=qa_policy_text
            )
            result = call_llm(prompt, temperature=0.3)
            if result and result.strip():
                cursor.execute(
                    "UPDATE manual_sections SET section_text = %s, updated_at = %s WHERE doc_id = %s AND section_name = %s",
                    (result.strip(), datetime.now().isoformat(), doc_id, section_name)
                )
                done_count += 1
        except Exception as e:
            print(f"[FILL_ALL] error for {section_name}: {e}")
            failed_sections.append({"section_name": section_name, "reason": str(e)})

    conn.commit()
    conn.close()
    return {
        "success": True, "total": total,
        "done_count": done_count, "failed_count": len(failed_sections),
        "failed_sections": failed_sections
    }


@router.post("/doc/{doc_id}/refine-all")
def refine_all(doc_id: str):
    """RAG ÏµúÏ†ÅÌôîÎßå Ï†Ñ ÏÑπÏÖòÏóê Ïã§Ìñâ."""
    if not is_llm_available():
        raise HTTPException(status_code=503, detail="LLM ÏÇ¨Ïö© Î∂àÍ∞Ä")

    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")
    raw_text = doc["raw_text"] or ""
    raw_text_safe = raw_text[:4000] if raw_text else "(ÏõêÎ≥∏ ÏóÜÏùå)"

    cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    all_sections = cursor.fetchall()
    if not all_sections:
        conn.close()
        raise HTTPException(status_code=400, detail="No sections found")

    total = len(all_sections)
    done_count = 0
    failed_sections = []

    for sec in all_sections:
        section_name = sec["section_name"]
        try:
            prompt = FINALIZE_SECTION_TEXT_PROMPT_V1.format(
                section_text=sec["section_text"],
                raw_text=raw_text_safe
            )
            result = call_llm(prompt, temperature=0.3)
            if result and result.strip():
                cursor.execute(
                    "UPDATE manual_sections SET section_text = %s, updated_at = %s WHERE doc_id = %s AND section_name = %s",
                    (result.strip(), datetime.now().isoformat(), doc_id, section_name)
                )
                done_count += 1
        except Exception as e:
            print(f"[REFINE_ALL] error for {section_name}: {e}")
            failed_sections.append({"section_name": section_name, "reason": str(e)})

    conn.commit()
    conn.close()
    return {
        "success": True, "total": total,
        "done_count": done_count, "failed_count": len(failed_sections),
        "failed_sections": failed_sections
    }


class PhaseUpdate(BaseModel):
    phase: str  # "fill", "refine", "gate"


@router.post("/doc/{doc_id}/completed-phase")
def save_completed_phase(doc_id: str, req: PhaseUpdate):
    """Add a completed phase to the document's completed_phases field."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT completed_phases FROM documents WHERE doc_id = %s", (doc_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")
    existing = row["completed_phases"] or ""
    phases = [p for p in existing.split(",") if p]
    if req.phase not in phases:
        phases.append(req.phase)
    cursor.execute("UPDATE documents SET completed_phases = %s WHERE doc_id = %s",
                   (",".join(phases), doc_id))
    conn.commit()
    conn.close()
    return {"success": True, "completed_phases": ",".join(phases)}


@router.post("/doc/{doc_id}/snapshot-ai-text")
def snapshot_ai_text(doc_id: str):
    """Copy current section_text to ai_text for all sections (after AI operation)."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE manual_sections SET ai_text = section_text WHERE doc_id = %s",
        (doc_id,)
    )
    updated = cursor.rowcount
    conn.commit()
    conn.close()
    return {"success": True, "updated": updated}


class SnapshotSectionAiText(BaseModel):
    section_name: str


@router.post("/doc/{doc_id}/snapshot-ai-text-section")
def snapshot_ai_text_section(doc_id: str, req: SnapshotSectionAiText):
    """Copy current section_text to ai_text for a single section."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE manual_sections SET ai_text = section_text WHERE doc_id = %s AND section_name = %s",
        (doc_id, req.section_name)
    )
    conn.commit()
    conn.close()
    return {"success": True}


@router.post("/doc/{doc_id}/section/{section_name}/gate-dismiss")
def gate_dismiss(doc_id: str, section_name: str, body: dict = Body(...)):
    """Toggle dismissed flag on a gate reason. Recalculate status from active reasons."""
    reason_index = body.get("reason_index")
    if reason_index is None:
        raise HTTPException(status_code=400, detail="reason_index required")
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT gate_reasons_json, gate_status FROM manual_sections WHERE doc_id = %s AND section_name = %s",
        (doc_id, section_name)
    )
    sec = cursor.fetchone()
    if not sec:
        conn.close()
        raise HTTPException(status_code=404, detail="Section not found")
    reasons = []
    try:
        reasons = json.loads(sec["gate_reasons_json"] or "[]")
    except Exception:
        pass
    if not (0 <= reason_index < len(reasons)):
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid reason_index")
    # Toggle dismissed flag
    reasons[reason_index]["dismissed"] = not reasons[reason_index].get("dismissed", False)
    # Recalculate status from active (non-dismissed) reasons only
    active = [r for r in reasons if not r.get("dismissed")]
    if not active:
        new_status = "PASS"
    else:
        has_high = any(r.get("severity") == "HIGH" for r in active)
        high_types = {r.get("type") for r in active if r.get("severity") == "HIGH"}
        block_types = {"PII_RISK", "CONFLICT", "OMISSION", "HALLUCINATION_RISK"}
        if has_high and high_types & block_types:
            new_status = "BLOCK"
        else:
            new_status = "NEED_FIX"
    cursor.execute(
        "UPDATE manual_sections SET gate_reasons_json = %s, gate_status = %s, gate_stale = 0, updated_at = %s WHERE doc_id = %s AND section_name = %s",
        (json.dumps(reasons, ensure_ascii=False), new_status, datetime.now().isoformat(), doc_id, section_name)
    )
    conn.commit()
    conn.close()
    return {"success": True, "status": new_status, "reasons": reasons}


@router.post("/doc/{doc_id}/section/{section_name}/gate-stale")
def set_gate_stale(doc_id: str, section_name: str):
    """Mark section as gate_stale (saved without gate re-check)."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "UPDATE manual_sections SET gate_stale = 1, updated_at = %s WHERE doc_id = %s AND section_name = %s",
        (datetime.now().isoformat(), doc_id, section_name)
    )
    conn.commit()
    conn.close()
    return {"success": True}


# ============ STEP 2: QUALITY GATE (document-level) ============

QUALITY_GATE_PROMPT = """ÎãπÏã†ÏùÄ ÏïÑÌååÌä∏ Ïö¥ÏòÅ Îß§Îâ¥ÏñºÏùò ÌíàÏßàÏùÑ Í≤ÄÏ¶ùÌïòÎäî Ï†ÑÎ¨∏Í∞ÄÏûÖÎãàÎã§.

ÏïÑÎûò Îß§Îâ¥Ïñº ÏÑπÏÖòÎì§ÏùÑ Í≤ÄÌÜ†ÌïòÍ≥† Ïù¥ÏäàÎ•º JSON Î∞∞Ïó¥Î°ú Î∞òÌôòÌïòÏÑ∏Ïöî.

Ïù¥Ïäà ÌÉÄÏûÖ:
- MISSING (RED): ÌôòÎ∂à/ÏòàÏïΩ/Ïö¥ÏòÅÏãúÍ∞Ñ/Í∂åÌïú Ï§ë ÌïµÏã¨ Ï†ïÎ≥¥Í∞Ä ÏôÑÏ†ÑÌûà ÏóÜÏùå
- AMBIGUOUS (YELLOW): "ÏÉÅÌô©Ïóê Îî∞Îùº", "Í∞ÄÎä•ÌïòÎ©¥", "Ï†ÅÎãπÌûà", "ÌòëÏùò ÌõÑ" Îì± Î™®Ìò∏Ìïú ÌëúÌòÑ
- CONFLICT (RED): Í∞ôÏùÄ Ï£ºÏ†úÏóêÏÑú ÏÉÅÎ∞òÎêú Í∑úÏπô Î∞úÍ≤¨
- PII_RISK (RED): Ï£ºÎØºÎ≤àÌò∏/Ï†ÑÌôîÎ≤àÌò∏ Îì± Í∞úÏù∏Ï†ïÎ≥¥ Ìå®ÌÑ¥
- API_NEEDED (YELLOW): "ÏòàÏïΩ ÏÉùÏÑ±", "Î¨∏Ïûê Î∞úÏÜ°", "Í∞ïÏ¢å Ï∂îÍ∞Ä" Îì± ÏãúÏä§ÌÖú Ïó∞Îèô ÌïÑÏöî

Í∞Å Ïù¥Ïäà ÌòïÏãù:
{{"severity": "RED|YELLOW", "issue_type": "ÌÉÄÏûÖ", "message": "ÏÑ§Î™Ö", "suggestion": "Ìï¥Í≤∞Î∞©Ïïà"}}

Îß§Îâ¥Ïñº ÎÇ¥Ïö©:
{sections_text}

JSON Î∞∞Ïó¥Îßå Î∞òÌôòÌïòÏÑ∏Ïöî."""


@router.post("/doc/{doc_id}/quality-gate")
def quality_gate(doc_id: str):
    """Run quality checks on manual sections."""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    sections = cursor.fetchall()
    if not sections:
        conn.close()
        raise HTTPException(status_code=400, detail="Manualize first")
    
    sections_text = "\n\n".join([f"[{s['section_name']}]\n{s['section_text']}" for s in sections])
    
    issues = []
    
    # Rule-based checks first
    # PII check
    pii_patterns = [
        (r'\d{6}-\d{7}', 'Ï£ºÎØºÎ≤àÌò∏'),
        (r'010-?\d{4}-?\d{4}', 'Ï†ÑÌôîÎ≤àÌò∏'),
    ]
    for pattern, pii_type in pii_patterns:
        if re.search(pattern, sections_text):
            issues.append({
                "severity": "RED",
                "issue_type": "PII_RISK",
                "message": f"{pii_type} Ìå®ÌÑ¥Ïù¥ Î∞úÍ≤¨ÎêòÏóàÏäµÎãàÎã§",
                "suggestion": "Í∞úÏù∏Ï†ïÎ≥¥Î•º ÏÇ≠Ï†úÌïòÍ±∞ÎÇò ÎßàÏä§ÌÇπÌïòÏÑ∏Ïöî"
            })
    
    # Ambiguous phrases check
    ambiguous_phrases = ["ÏÉÅÌô©Ïóê Îî∞Îùº", "Í∞ÄÎä•ÌïòÎ©¥", "Ï†ÅÎãπÌûà", "ÌòëÏùò ÌõÑ", "Í≤ΩÏö∞Ïóê Îî∞Îùº", "ÌïÑÏöîÏãú"]
    for phrase in ambiguous_phrases:
        if phrase in sections_text:
            issues.append({
                "severity": "YELLOW",
                "issue_type": "AMBIGUOUS",
                "message": f"Î™®Ìò∏Ìïú ÌëúÌòÑ Î∞úÍ≤¨: '{phrase}'",
                "suggestion": "Íµ¨Ï≤¥Ï†ÅÏù∏ Í∏∞Ï§ÄÏù¥ÎÇò Ï°∞Í±¥ÏúºÎ°ú Î™ÖÏãúÌïòÏÑ∏Ïöî"
            })
    
    # API needed check
    api_phrases = ["ÏòàÏïΩ ÏÉùÏÑ±", "ÏòàÏïΩ Ï∑®ÏÜå", "Î¨∏Ïûê Î∞úÏÜ°", "SMS", "Í∞ïÏ¢å Ï∂îÍ∞Ä", "Í∞ïÏ¢å ÏÇ≠Ï†ú", "ÌöåÏõê Îì±Î°ù"]
    for phrase in api_phrases:
        if phrase in sections_text:
            issues.append({
                "severity": "YELLOW",
                "issue_type": "API_NEEDED",
                "message": f"ÏãúÏä§ÌÖú Ïó∞Îèô ÌïÑÏöî: '{phrase}'",
                "suggestion": "Ìï¥Îãπ Í∏∞Îä•Ïùò API Ïä§ÌéôÏùÑ Ï†ïÏùòÌïòÏÑ∏Ïöî"
            })
    
    # Missing check for critical sections
    for s in sections:
        if s["section_text"] == "Ï†ïÎ≥¥ ÏóÜÏùå" and s["section_name"] in ["ÌôòÎ∂à/ÏúÑÏïΩ/Ï†ïÏÇ∞", "ÏòàÏïΩ/Ï∑®ÏÜå/Î≥ÄÍ≤Ω", "Ïö¥ÏòÅÏãúÍ∞Ñ/Ìú¥Î¨¥"]:
            issues.append({
                "severity": "RED",
                "issue_type": "MISSING",
                "message": f"ÌïÑÏàò ÏÑπÏÖò '{s['section_name']}'Ïùò ÎÇ¥Ïö©Ïù¥ ÏóÜÏäµÎãàÎã§",
                "suggestion": "Ìï¥Îãπ Í∑úÏ†ïÏùÑ Ï∂îÍ∞ÄÌïòÏÑ∏Ïöî"
            })
    
    # LLM-based additional checks if available
    llm_error_msg = None
    if is_llm_available() and len(issues) < 5:
        try:
            content = call_llm(QUALITY_GATE_PROMPT.format(sections_text=sections_text[:6000]), temperature=0.3)
            if content:
                json_match = re.search(r'\[[\s\S]*\]', content)
                if json_match:
                    llm_issues = _clean_llm_json(json_match.group())
                    issues.extend(llm_issues[:5])  # Limit LLM issues
            else:
                llm_error_msg = "LLM ÏùëÎãµÏù¥ ÎπÑÏñ¥ÏûàÏäµÎãàÎã§. (API Ìï†ÎãπÎüâ Ï¥àÍ≥º Í∞ÄÎä•ÏÑ±)"
        except Exception as e:
            llm_error_msg = f"LLM Ìò∏Ï∂ú Ïã§Ìå®: {str(e)}"
            print(f"[QUALITY_GATE] LLM error: {e}")
    
    # Clear old issues and save new
    cursor.execute("DELETE FROM qa_issues WHERE doc_id = %s", (doc_id,))
    for issue in issues:
        issue_id = f"issue_{uuid.uuid4().hex[:8]}"
        cursor.execute(
            "INSERT INTO qa_issues (issue_id, doc_id, severity, issue_type, message, suggestion, status) VALUES (%s, %s, %s, %s, %s, %s, 'OPEN')",
            (issue_id, doc_id, issue.get("severity", "YELLOW"), issue.get("issue_type", "OTHER"), 
             issue.get("message", ""), issue.get("suggestion", ""))
        )
    
    conn.commit()
    conn.close()
    
    red_count = len([i for i in issues if i.get("severity") == "RED"])
    yellow_count = len([i for i in issues if i.get("severity") == "YELLOW"])
    
    return {
        "success": True, 
        "doc_id": doc_id, 
        "red_count": red_count, 
        "yellow_count": yellow_count, 
        "issues": issues,
        "llm_error": llm_error_msg,
        "api_specs": extract_api_spec(doc_id).get("specs", [])
    }


# ============ STEP 2: UPDATE SECTIONS ============

class SectionsUpdate(BaseModel):
    sections: dict # {section_name: text}

@router.put("/doc/{doc_id}/sections")
def update_sections(doc_id: str, req: SectionsUpdate):
    """Update manual sections manually."""
    conn = get_connection()
    cursor = conn.cursor()

    for name, text in req.sections.items():
        cursor.execute(
            "UPDATE manual_sections SET section_text = %s WHERE doc_id = %s AND section_name = %s",
            (text, doc_id, name)
        )
    
    cursor.execute("UPDATE documents SET updated_at = %s WHERE doc_id = %s", (datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()
    
    return {"success": True, "doc_id": doc_id}


# ============ STEP 2: AI HELPER (Fill/Refine) ============

class RefineRequest(BaseModel):
    text: str
    task: str  # "refine", "fill", "recommend"
    context: Optional[str] = None  # Section name or issue message
    allow_qa: Optional[str] = None  # "true" or "false" (for fill task)


def _clean_llm_json(raw: str):
    """Clean and parse JSON from LLM response. Handles common issues.
    Returns dict or list depending on input."""
    raw = re.sub(r',\s*([}\]])', r'\1', raw)       # trailing comma
    raw = raw.replace('\n', ' ').replace('\r', '')   # newlines
    raw = re.sub(r'[\x00-\x1f]', ' ', raw)          # control chars
    # Attempt 1: direct parse
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    # Attempt 2: single ‚Üí double quotes
    try:
        return json.loads(raw.replace("'", '"'))
    except json.JSONDecodeError:
        pass
    # Attempt 3: fix unescaped quotes inside string values
    # e.g. "message": "Î¨∏Ï†úÎäî "Ïù¥Í≤É" ÏûÖÎãàÎã§" ‚Üí "message": "Î¨∏Ï†úÎäî \"Ïù¥Í≤É\" ÏûÖÎãàÎã§"
    try:
        fixed = re.sub(
            r'(?<=: )"((?:[^"\\]|\\.)*?)"(?=\s*[,}\]])',
            lambda m: '"' + m.group(1) + '"',
            raw
        )
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass
    # Attempt 4: try closing at each } or ] from the end (handles trailing garbage)
    close_positions = [i for i, c in enumerate(raw) if c in ('}', ']')]
    for pos in reversed(close_positions):
        try:
            return json.loads(raw[:pos + 1])
        except json.JSONDecodeError:
            continue
    raise json.JSONDecodeError("_clean_llm_json: all attempts failed", raw, 0)


def _flatten_manualize_json(parsed: dict) -> dict:
    """Convert MANUALIZE_PROMPT JSON output to flat {section_name: section_text} dict.
    Formats each section as structured markdown (## / ### / - bullets)."""
    sections_map = {}
    for sec in parsed.get("sections", []):
        name = sec.get("name", sec.get("section_id", "general"))
        lines = [f"## {name}"]
        for rule in sec.get("content", []):
            title = rule.get("title", "")
            if title:
                lines.append(f"### {title}")
            for bullet in rule.get("bullets", []):
                lines.append(f"- {bullet}")
            # structured fields
            st = rule.get("structured", {})
            for key in ("target", "condition", "owner", "channel"):
                val = st.get(key, "")
                if val:
                    lines.append(f"- {key}: {val}")
            for proc_item in st.get("procedure", []):
                lines.append(f"- {proc_item}")
            for exc_item in st.get("exceptions", []):
                lines.append(f"- ÏòàÏô∏: {exc_item}")
        section_text = "\n".join(lines)
        sections_map[name] = section_text
    # Fallback: if no sections parsed, try summary
    if not sections_map:
        summary = parsed.get("summary", "")
        if summary:
            sections_map["general"] = f"## ÏöîÏïΩ\n- {summary}"
        else:
            sections_map["general"] = "Ï†ïÎ≥¥ ÏóÜÏùå"
    return sections_map


def _to_bool_allow_qa(val) -> bool:
    """Normalize allow_qa to bool. Accepts bool, str, int, None."""
    if val is None:
        return False
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.strip().lower() in ("true", "1", "yes")
    if isinstance(val, (int, float)):
        return bool(val)
    return False


# ============ CHUNK-BASED PIPELINE: PROMPTS ============

SPLIT_PROMPT = """ÎãπÏã†ÏùÄ Í∏¥ Î¨∏ÏÑúÎ•º RAG Ï≤òÎ¶¨Ïóê Ï†ÅÌï©Ìïú chunk(Ï≤≠ÌÅ¨)Î°ú Î∂ÑÌï†ÌïòÎäî Ï†ÑÎ¨∏Í∞ÄÏûÖÎãàÎã§.

[ÏûÖÎ†•]
raw_text: {raw_text}

[Í∑úÏπô]
1) Ìó§Îî©(#, ##, Ïà´Ïûê. Îì±) Í≤ΩÍ≥ÑÎ•º Ïö∞ÏÑ†ÏúºÎ°ú Î∂ÑÌï†Ìï©ÎãàÎã§.
2) Í∞Å chunkÎäî 3,000~8,000Ïûê Î™©Ìëú, ÏµúÎåÄ 12,000ÏûêÎ•º ÎÑòÏßÄ ÏïäÏäµÎãàÎã§.
3) 3,000Ïûê ÎØ∏ÎßåÏùò ÏßßÏùÄ ÏÑπÏÖòÏùÄ Ïù∏Ï†ë ÏÑπÏÖòÍ≥º Ìï©Ïπ©ÎãàÎã§.
4) start_anchor, end_anchorÎäî ÏõêÎ¨∏ÏóêÏÑú Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨Ìïú Î¨∏ÏûêÏó¥(20~60Ïûê)ÏûÖÎãàÎã§.
   - start_anchor: Ìï¥Îãπ chunkÍ∞Ä ÏãúÏûëÌïòÎäî ÏõêÎ¨∏ Î¨∏Ïû•/Ìó§Îî©Ïùò Ï≤´ Î∂ÄÎ∂Ñ
   - end_anchor: Ìï¥Îãπ chunkÍ∞Ä ÎÅùÎÇòÎäî ÏõêÎ¨∏ Î¨∏Ïû•Ïùò ÎßàÏßÄÎßâ Î∂ÄÎ∂Ñ
5) split_basis: Î∂ÑÌï† Í∑ºÍ±∞ (Ïòà: "heading", "length", "topic_shift")
6) chunk_idÎäî "C1", "C2"... ÏàúÏÑúÏûÖÎãàÎã§.

[Ï∂úÎ†• ÌòïÏãù]
RFC8259 JSON Î∞∞Ïó¥Îßå Ï∂úÎ†•. ÏΩîÎìúÎ∏îÎ°ù/ÏÑ§Î™Ö/Ï£ºÏÑù Í∏àÏßÄ.

[
  {{
    "chunk_id": "C1",
    "start_anchor": "ÏõêÎ¨∏ ÏãúÏûë Î∂ÄÎ∂Ñ ÌÖçÏä§Ìä∏",
    "end_anchor": "ÏõêÎ¨∏ ÎÅù Î∂ÄÎ∂Ñ ÌÖçÏä§Ìä∏",
    "split_basis": "heading",
    "notes": ""
  }}
]

[Ï†úÏïΩ]
- anchorÎäî Î∞òÎìúÏãú ÏõêÎ¨∏(raw_text)Ïóê Ï°¥Ïû¨ÌïòÎäî Î¨∏ÏûêÏó¥Ïù¥Ïñ¥Ïïº Ìï©ÎãàÎã§.
- chunk Í∞Ñ Í≤πÏπ® ÏóÜÏù¥ Ï†ÑÏ≤¥ Î¨∏ÏÑúÎ•º Îπ†ÏßêÏóÜÏù¥ Ïª§Î≤ÑÌï¥Ïïº Ìï©ÎãàÎã§.
- JSON Î∞∞Ïó¥ Ïô∏ Ïñ¥Îñ§ Ï∂úÎ†•ÎèÑ Í∏àÏßÄÌï©ÎãàÎã§."""

MANUALIZE_CHUNK_PROMPT = """ÎãπÏã†ÏùÄ ÏòÅÏóÖ/Ïö¥ÏòÅ Î¨∏ÏÑúÏùò ÏùºÎ∂Ä(chunk)Î•º RAGÏóê ÎÑ£Í∏∞ ÏúÑÌïú
"Íµ¨Ï°∞ÌôîÎêú Ï†ïÎ≥¥ Ï∂îÏ∂úÍ∏∞"ÏûÖÎãàÎã§.

‚ö†Ô∏è Ïù¥ ÏûëÏóÖÏùÄ ÏöîÏïΩÏù¥ ÏïÑÎãôÎãàÎã§. Î¨∏ÏÑúÎ•º Ï§ÑÏù¥Í±∞ÎÇò ÏïïÏ∂ïÌïòÎäî ÏûëÏóÖÏù¥ ÏïÑÎãôÎãàÎã§.

[ÏµúÏö∞ÏÑ† Î™©Ìëú]
- ÏûÖÎ†•Îêú chunk(raw_chunk)Ïùò Ï†ïÎ≥¥ÏôÄ Íµ¨Ï°∞Î•º ÏµúÎåÄÌïú Í∑∏ÎåÄÎ°ú Î≥¥Ï°¥ÌïòÏó¨ Íµ¨Ï°∞ÌôîÌï©ÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏóÜÎäî Ï†ïÎ≥¥Î•º Ï∂îÍ∞Ä/Ï∂îÏ∏°/Ï∞ΩÏûëÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏûàÎäî Ï†ïÎ≥¥Î•º ÏÇ≠Ï†úÌïòÍ±∞ÎÇò ÏÉùÎûµÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- Í∞úÏù∏Ï†ïÎ≥¥(Ï†ÑÌôîÎ≤àÌò∏, Ïù¥Î©îÏùº, Í≥ÑÏ¢åÎ≤àÌò∏ Îì±)Îäî ***Î°ú ÎßàÏä§ÌÇπÌï©ÎãàÎã§.

[ÏûÖÎ†•]
raw_chunk: {raw_chunk}

[Ï∂úÎ†• ÌòïÏãù]
RFC8259 Ïú†Ìö® JSONÎßå Ï∂úÎ†•. ÏΩîÎìúÎ∏îÎ°ù, ÏÑ§Î™Ö, Ï£ºÏÑù, ÎßàÌÅ¨Îã§Ïö¥ Í∏àÏßÄ.

{{
  "section_name": "Ïù¥ chunkÏùò Ï£ºÏ†úÎ•º ÎåÄÌëúÌïòÎäî ÏÑπÏÖòÎ™Ö",
  "section_text": "Íµ¨Ï°∞ÌôîÎêú ÌÖçÏä§Ìä∏ (## Ìó§Îî©, ### Ìï≠Î™©, - Î∂àÎ¶ø ÌòïÏãù)",
  "evidence_spans": [
    {{
      "span_text": "raw_chunkÏóêÏÑú Î≥µÏÇ¨Ìïú Í∑ºÍ±∞ Î¨∏Ïû• (20~80Ïûê)",
      "maps_to": "section_text ÎÇ¥ ÎåÄÏùë ÏúÑÏπò ÏÑ§Î™Ö (Ìï≠Î™©Î™Ö ÎòêÎäî Î∂àÎ¶ø ÏöîÏïΩ)",
      "is_pii": false
    }}
  ]
}}

[evidence_spans Í∑úÏπô]
- ÏµúÏÜå 2Í∞ú Ïù¥ÏÉÅ ÏûëÏÑ± (Í∞ÄÎä•ÌïòÎ©¥ ÌïµÏã¨ Ï†ïÎ≥¥ÎßàÎã§ 1Í∞ú)
- span_textÎäî raw_chunk ÏõêÎ¨∏ÏóêÏÑú Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨ (verbatim)
- is_pii=trueÏù∏ spanÏùÄ ÎßàÏä§ÌÇπÎêú ÌòïÌÉúÎ°ú ÏûëÏÑ±
- maps_toÎäî section_textÏùò Ïñ¥Îäê Î∂ÄÎ∂ÑÏóê ÎåÄÏùëÌïòÎäîÏßÄ Í∞ÑÎã®Ìûà ÌëúÏãú

[section_text ÏûëÏÑ± Í∑úÏπô]
- "## ÏÑπÏÖòÎ™Ö"ÏúºÎ°ú ÏãúÏûë
- "### Ìï≠Î™©Ï†úÎ™©"ÏúºÎ°ú ÌïòÏúÑ Íµ¨Î∂Ñ
- "- " Î∂àÎ¶øÏúºÎ°ú Î≥∏Î¨∏
- ÏõêÎ¨∏ Ï†ïÎ≥¥ ÏÇ≠Ï†ú/ÏöîÏïΩ/ÏïïÏ∂ï Í∏àÏßÄ
- ÏõêÎ¨∏ Î∂àÎ¶ø Í∞úÏàòÏôÄ Ìï≠Î™© Í∑∏ÎåÄÎ°ú Î≥¥Ï°¥

[Í∏àÏßÄ]
- Ï∂îÏ∏°, ÏùºÎ∞ò ÏÉÅÏãù Î≥¥ÏôÑ, ÏÉà Ï†ïÎ≥¥ Ï∂îÍ∞Ä
- JSON Ïô∏ Ï∂úÎ†•"""


# ============ CHUNK-BASED PIPELINE: HELPERS ============

def _resolve_anchors(raw_text: str, chunks_raw: list) -> list:
    """Resolve start_anchor/end_anchor to char_start/char_end positions in raw_text."""
    resolved = []
    search_from = 0
    for i, chunk in enumerate(chunks_raw):
        start_anchor = chunk.get("start_anchor", "")
        end_anchor = chunk.get("end_anchor", "")

        # Find start position
        char_start = raw_text.find(start_anchor, search_from) if start_anchor else search_from
        if char_start < 0:
            # Fallback: try from beginning
            char_start = raw_text.find(start_anchor) if start_anchor else search_from
        if char_start < 0:
            char_start = search_from

        # Find end position
        if end_anchor:
            end_search = max(char_start, search_from)
            end_pos = raw_text.find(end_anchor, end_search)
            if end_pos >= 0:
                char_end = end_pos + len(end_anchor)
            else:
                # Fallback: next chunk's start or end of text
                char_end = len(raw_text) if i == len(chunks_raw) - 1 else None
        else:
            char_end = len(raw_text) if i == len(chunks_raw) - 1 else None

        # Deferred end for non-last chunks
        if char_end is None:
            char_end = len(raw_text)  # will be adjusted by next chunk

        resolved.append({
            **chunk,
            "char_start": char_start,
            "char_end": char_end,
        })
        search_from = char_start + 1

    # Adjust: ensure no gaps/overlaps between consecutive chunks
    for i in range(len(resolved) - 1):
        if resolved[i + 1]["char_start"] < resolved[i]["char_end"]:
            resolved[i]["char_end"] = resolved[i + 1]["char_start"]

    # Extract raw_chunk text
    for r in resolved:
        r["raw_chunk"] = raw_text[r["char_start"]:r["char_end"]]

    return resolved


def _save_doc_chunks(doc_id: str, chunks: list, cursor) -> list:
    """Save resolved chunks to doc_chunks table. Returns list with chunk_ids."""
    cursor.execute("DELETE FROM doc_chunks WHERE doc_id = %s", (doc_id,))
    saved = []
    for i, chunk in enumerate(chunks):
        chunk_id = f"dchunk_{uuid.uuid4().hex[:8]}"
        cursor.execute("""
            INSERT INTO doc_chunks (chunk_id, doc_id, chunk_index, start_anchor, end_anchor,
                                    char_start, char_end, split_basis, notes, raw_chunk)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (chunk_id, doc_id, i,
              chunk.get("start_anchor", ""), chunk.get("end_anchor", ""),
              chunk.get("char_start", 0), chunk.get("char_end", 0),
              chunk.get("split_basis", ""), chunk.get("notes", ""),
              chunk.get("raw_chunk", "")))
        saved.append({**chunk, "chunk_id": chunk_id, "chunk_index": i})
    return saved


def _index_evidence_spans(spans: list, raw_chunk: str) -> list:
    """Index span_text positions within raw_chunk. Sets char_start/char_end or -1 if not found."""
    indexed = []
    for span in spans:
        span_text = span.get("span_text", "")
        is_pii = span.get("is_pii", False)

        if is_pii or not span_text:
            indexed.append({**span, "char_start": -1, "char_end": -1})
            continue

        pos = raw_chunk.find(span_text)
        if pos < 0:
            # Fallback: case-insensitive
            pos = raw_chunk.lower().find(span_text.lower())
        if pos >= 0:
            indexed.append({**span, "char_start": pos, "char_end": pos + len(span_text)})
        else:
            indexed.append({**span, "char_start": -1, "char_end": -1})
    return indexed


def _conservative_dedup(sections: list) -> list:
    """Remove exact-duplicate sections by section_name. Keep first occurrence."""
    seen = set()
    result = []
    for sec in sections:
        name = sec.get("section_name", "")
        if name not in seen:
            seen.add(name)
            result.append(sec)
    return result


# ============ BATCH MANUALIZE ============

MAX_BATCH_CHARS = 12000

MANUALIZE_BATCH_PROMPT = """ÎãπÏã†ÏùÄ 'Manualize Î∞∞Ïπò Ïã§ÌñâÍ∏∞'ÏûÖÎãàÎã§.
Ï§ëÏöî: ÏïÑÎûò [MANUALIZE_RULES]Ïùò Í∑úÏπôÏùÑ **Í∑∏ÎåÄÎ°ú Ï§ÄÏàò**ÌïòÏó¨ ÏûëÏóÖÌïòÏÑ∏Ïöî.

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[MANUALIZE_RULES - SOURCE OF TRUTH]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
- ÏõêÎ¨∏(section_text)Ïùò Ï†ïÎ≥¥ÏôÄ Íµ¨Ï°∞Î•º ÏµúÎåÄÌïú Í∑∏ÎåÄÎ°ú Î≥¥Ï°¥ÌïòÏó¨ Íµ¨Ï°∞ÌôîÌï©ÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏóÜÎäî Ï†ïÎ≥¥Î•º Ï∂îÍ∞Ä/Ï∂îÏ∏°/Ï∞ΩÏûëÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- ÏõêÎ¨∏Ïóê ÏûàÎäî Ï†ïÎ≥¥Î•º ÏÇ≠Ï†úÌïòÍ±∞ÎÇò ÏÉùÎûµÌïòÏßÄ ÏïäÏäµÎãàÎã§.
- ÏöîÏïΩ/ÏïïÏ∂ï/ÏùºÎ∞òÌôî Í∏àÏßÄ.
- Í∞úÏù∏Ï†ïÎ≥¥(Ï†ÑÌôîÎ≤àÌò∏, Ïù¥Î©îÏùº, Í≥ÑÏ¢åÎ≤àÌò∏ Îì±)Îäî ***Î°ú ÎßàÏä§ÌÇπÌï©ÎãàÎã§.
- section_text ÎÇ¥Î∂ÄÏùò Ìó§Îî©/Î≤àÌò∏/Íµ¨Î∂Ñ Íµ¨Ï°∞Î•º Í∑∏ÎåÄÎ°ú Îß§ÌïëÌïòÏÑ∏Ïöî.
- Íµ¨Î∂ÑÏù¥ ÏóÜÏúºÎ©¥ sections=1Í∞ú(general)Î°ú Ï≤òÎ¶¨ÌïòÏÑ∏Ïöî.
- Ï∂úÎ†• ÌòïÏãù: "## ÏÑπÏÖòÎ™Ö", "### Ìï≠Î™©Ï†úÎ™©", "- " Î∂àÎ¶ø.
- evidence_spans: ÏµúÏÜå 2Í∞ú, span_textÎäî ÏõêÎ¨∏ÏóêÏÑú Í∑∏ÎåÄÎ°ú Î≥µÏÇ¨(verbatim).

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[ÏûÖÎ†•]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
{batches_json}

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[ÏûëÏóÖ Í∑úÏπô - Î∞∞Ïπò Ï†ÑÏö©]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
1) item(=ÏÑπÏÖò) Í∞Ñ Ï†ïÎ≥¥ ÏÑûÍ∏∞ Í∏àÏßÄ. Í∞Å ÏÑπÏÖòÏùÄ ÎèÖÎ¶ΩÏ†ÅÏúºÎ°ú Manualize.
2) ÏÑπÏÖò Íµ¨Ï°∞ Î≥¥Ï°¥. ÏÑπÏÖò ÏàòÎ•º ÏûÑÏùò Ï°∞Ï†ï Í∏àÏßÄ.
3) oversize_sectionsÏóê ÏûàÎäî Ìï≠Î™©ÏùÄ Manualize ÌïòÏßÄ ÎßêÍ≥† errorsÏóê Í∏∞Î°ùÎßå.

‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
[Ï∂úÎ†• ÌòïÏãù]
‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
RFC8259 Ïú†Ìö® JSONÎßå Ï∂úÎ†•. ÏΩîÎìúÎ∏îÎ°ù/ÏÑ§Î™Ö/Ï£ºÏÑù Í∏àÏßÄ.

{{
  "batch_id": "<ÏûÖÎ†• batch_id>",
  "results": [
    {{
      "item_id": "S001",
      "section_name": "ÏÑπÏÖò Ï£ºÏ†úÎ•º ÎåÄÌëúÌïòÎäî Ïù¥Î¶Ñ",
      "section_text": "Íµ¨Ï°∞ÌôîÎêú ÌÖçÏä§Ìä∏ (## / ### / - ÌòïÏãù)",
      "evidence_spans": [
        {{
          "span_text": "ÏõêÎ¨∏ÏóêÏÑú Î≥µÏÇ¨Ìïú Í∑ºÍ±∞ Î¨∏Ïû• (20~80Ïûê)",
          "maps_to": "section_text ÎÇ¥ ÎåÄÏùë ÏúÑÏπò",
          "is_pii": false
        }}
      ]
    }}
  ],
  "errors": []
}}"""


def _build_batches(doc_chunks: list, doc_id: str, max_chars: int = MAX_BATCH_CHARS) -> list:
    """Group doc_chunks into batches. Section boundary = batch boundary.
    A section never splits across batches. If adding a section exceeds max_chars,
    it goes to the next batch."""
    batches = []
    current_items = []
    current_chars = 0
    oversize = []
    batch_index = 0

    for chunk in doc_chunks:
        raw_chunk = chunk.get("raw_chunk", "")
        char_len = len(raw_chunk)
        heading = chunk.get("notes", "") or ""
        if not heading and raw_chunk:
            first_line = raw_chunk.split("\n", 1)[0].strip().lstrip("#").strip()
            heading = first_line[:60] if first_line else f"Chunk {chunk.get('chunk_index', 0) + 1}"

        item = {
            "item_id": chunk["chunk_id"],
            "section_title": heading,
            "section_text": raw_chunk,
            "char_len": char_len,
        }

        # Oversize: single section exceeds max_chars
        if char_len > max_chars:
            oversize.append({
                "item_id": chunk["chunk_id"],
                "section_title": heading,
                "char_len": char_len,
                "reason": "section_text exceeds max_batch_chars"
            })
            continue

        # Would exceed limit ‚Üí close current batch, start new one
        if current_items and current_chars + char_len > max_chars:
            batch_index += 1
            batches.append({
                "doc_id": doc_id,
                "batch_id": f"B{batch_index:03d}",
                "batch_index": batch_index - 1,
                "max_batch_chars": max_chars,
                "items": current_items,
                "oversize_sections": [],
                "chunk_ids": [it["item_id"] for it in current_items],
            })
            current_items = []
            current_chars = 0

        current_items.append(item)
        current_chars += char_len

    # Last batch
    if current_items:
        batch_index += 1
        batches.append({
            "doc_id": doc_id,
            "batch_id": f"B{batch_index:03d}",
            "batch_index": batch_index - 1,
            "max_batch_chars": max_chars,
            "items": current_items,
            "oversize_sections": [],
            "chunk_ids": [it["item_id"] for it in current_items],
        })

    # Attach oversize to first batch (or create one)
    if oversize:
        if batches:
            batches[0]["oversize_sections"] = oversize
        else:
            batches.append({
                "doc_id": doc_id,
                "batch_id": "B001",
                "batch_index": 0,
                "max_batch_chars": max_chars,
                "items": [],
                "oversize_sections": oversize,
                "chunk_ids": [],
            })

    return batches


def _manualize_batch(batch: dict, doc_id: str) -> list:
    """Process a single batch via LLM. Returns list of per-item results."""
    items = batch.get("items", [])
    if not items:
        return []

    # Build batch JSON for prompt
    batch_input = {
        "doc_id": doc_id,
        "batch_id": batch["batch_id"],
        "max_batch_chars": batch["max_batch_chars"],
        "items": [
            {"item_id": it["item_id"], "section_title": it["section_title"],
             "section_text": it["section_text"], "char_len": it["char_len"]}
            for it in items
        ],
        "oversize_sections": batch.get("oversize_sections", []),
    }

    for attempt in range(2):
        try:
            content = call_llm(
                MANUALIZE_BATCH_PROMPT.format(
                    batches_json=json.dumps(batch_input, ensure_ascii=False)
                ),
                temperature=0.3
            )
            if content:
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    parsed = _clean_llm_json(json_match.group())
                    results = parsed.get("results", [])
                    if results:
                        return results
        except Exception as e:
            print(f"[MANUALIZE_BATCH] attempt {attempt + 1} failed for {batch['batch_id']}: {e}")

    # Fallback: return raw text as-is for each item
    fallback = []
    for it in items:
        fallback.append({
            "item_id": it["item_id"],
            "section_name": it["section_title"] or f"[ÎØ∏Ï≤òÎ¶¨] {it['item_id']}",
            "section_text": it["section_text"][:3000],
            "evidence_spans": [],
        })
    return fallback


# ============ CHUNK-BASED PIPELINE: CORE FUNCTIONS ============

def _split_document(raw_text: str, doc_id: str, cursor):
    """Split document into chunks using rule-based splitting (no LLM).
    Returns list of resolved chunks with raw_chunk text."""
    chunks_raw = []

    # Rule-based: heading split
    heading_chunks = _split_raw_by_headings(raw_text)
    if heading_chunks:
        for i, hc in enumerate(heading_chunks):
            start_anchor = hc["body"][:50].strip() if hc["body"] else ""
            end_anchor = hc["body"][-50:].strip() if hc["body"] else ""
            chunks_raw.append({
                "chunk_id": f"C{i + 1}",
                "start_anchor": start_anchor,
                "end_anchor": end_anchor,
                "split_basis": "heading",
                "notes": hc["heading"]
            })

    # Fallback: single chunk (entire document)
    if not chunks_raw:
        print(f"[SPLIT] No headings found, single chunk for {doc_id}")
        chunks_raw = [{
            "chunk_id": "C1",
            "start_anchor": raw_text[:50].strip(),
            "end_anchor": raw_text[-50:].strip(),
            "split_basis": "single_document",
            "notes": "No split possible"
        }]

    # Resolve anchors to char positions and extract raw_chunk
    resolved = _resolve_anchors(raw_text, chunks_raw)

    # Save to DB
    saved = _save_doc_chunks(doc_id, resolved, cursor)
    print(f"[SPLIT] {doc_id}: {len(saved)} chunks created")
    return saved


def _manualize_chunk(chunk: dict, doc_id: str) -> dict:
    """Manualize a single chunk. Returns dict with section_name, section_text, evidence_spans, chunk_id.
    Retry once on failure."""
    raw_chunk = chunk.get("raw_chunk", "")
    chunk_id = chunk.get("chunk_id", "")

    for attempt in range(2):
        try:
            content = call_llm(
                MANUALIZE_CHUNK_PROMPT.format(raw_chunk=raw_chunk[:10000]),
                temperature=0.3
            )
            if content:
                json_match = re.search(r'\{[\s\S]*\}', content)
                if json_match:
                    result = _clean_llm_json(json_match.group())
                    section_name = result.get("section_name", f"Chunk-{chunk.get('chunk_index', 0)}")
                    section_text = result.get("section_text", "")
                    evidence_spans = result.get("evidence_spans", [])

                    # Index evidence spans
                    indexed_spans = _index_evidence_spans(evidence_spans, raw_chunk)

                    return {
                        "section_name": section_name,
                        "section_text": section_text,
                        "evidence_spans": indexed_spans,
                        "chunk_id": chunk_id,
                        "success": True
                    }
        except Exception as e:
            print(f"[MANUALIZE_CHUNK] attempt {attempt + 1} failed for {chunk_id}: {e}")

    # Both attempts failed ‚Äî return raw chunk as-is
    return {
        "section_name": f"[ÎØ∏Ï≤òÎ¶¨] Chunk-{chunk.get('chunk_index', 0)}",
        "section_text": raw_chunk[:3000],
        "evidence_spans": [],
        "chunk_id": chunk_id,
        "success": False
    }


def _gate_chunk(section_text: str, raw_chunk: str) -> dict:
    """Run gate check on a chunk using only the raw_chunk (not full raw_text)."""
    gate_result = {"status": "PASS", "score": 100, "reasons": [], "required_actions": []}

    if not is_llm_available():
        return gate_result

    try:
        content = call_llm(GATE_CHECK_PROMPT.format(
            section_text=section_text[:3000],
            raw_text=raw_chunk[:4000]
        ), temperature=0.3)
        if content:
            gm = re.search(r'\{[\s\S]*\}', content)
            if gm:
                gate_result = _clean_llm_json(gm.group())
    except Exception as e:
        print(f"[GATE_CHUNK] error: {e}")

    return gate_result


def _merge_chunks(chunk_results: list, gate_results: list) -> dict:
    """Merge chunk manualize results. PASS ‚Üí BODY, NEED_FIX ‚Üí APPENDIX.
    Returns {sections_map, gate_map, evidence_map, merge_info}."""
    sections_map = {}
    gate_map = {}
    evidence_map = {}
    merge_info = {}

    for i, result in enumerate(chunk_results):
        gate = gate_results[i] if i < len(gate_results) else {"status": "PASS", "score": 100, "reasons": []}
        name = result["section_name"]
        gate_status = gate.get("status", "PASS")

        # Determine merge_status
        if gate_status in ("NEED_FIX", "BLOCK"):
            merge_status = "APPENDIX"
            display_name = f"[Î∂ÄÎ°ù] {name}"
        else:
            merge_status = "BODY"
            display_name = name

        # Conservative dedup: if same name exists, append index
        if display_name in sections_map:
            display_name = f"{display_name} ({i + 1})"

        sections_map[display_name] = result["section_text"]
        gate_map[display_name] = gate
        evidence_map[display_name] = result.get("evidence_spans", [])
        merge_info[display_name] = {
            "merge_status": merge_status,
            "chunk_id": result.get("chunk_id", ""),
            "original_name": name,
        }

    return {
        "sections_map": sections_map,
        "gate_map": gate_map,
        "evidence_map": evidence_map,
        "merge_info": merge_info,
    }


FILL_SECTION_TEXT_PROMPT_V3 = """ÎãπÏã†ÏùÄ RAG Ï≤≠ÌÅ¨(ÏÑπÏÖò ÌÖçÏä§Ìä∏)Î•º 'ÎèÖÎ¶ΩÏ†ÅÏúºÎ°ú Ïù¥Ìï¥ Í∞ÄÎä•Ìïú Îß§Îâ¥Ïñº'Î°ú Î≥¥Í∞ïÌïòÎäî Ìé∏ÏßëÏûêÏûÖÎãàÎã§.
Ï§ëÏöî: Ïù¥ ÏûëÏóÖÏùÄ 'ÏÉà Ï†ïÎ≥¥ Ï∂îÍ∞Ä'Í∞Ä ÏïÑÎãàÎùº, ÎèôÏùº Î¨∏ÏÑú(raw_text) ÎÇ¥Î∂ÄÏùò Í¥ÄÎ†® ÎÇ¥Ïö©ÏùÑ Î™®ÏïÑ Ïû¨Íµ¨ÏÑ±ÌïòÎäî Í≤ÉÏûÖÎãàÎã§.

[ÏûÖÎ†•]
section_text: {section_text}
raw_text: {raw_text}

[Ï†àÎåÄ Í∑úÏπô]
1) ÏõêÎ¨∏Ïóê ÏóÜÎäî Ï†ïÎ≥¥(ÏàòÏπò/Í∏∞Í∞Ñ/Í∏àÏï°/Ï†ïÏ±Ö/ÏòàÏô∏)Î•º Ï†àÎåÄ ÎßåÎì§ÏßÄ ÎßàÏÑ∏Ïöî.
2) Í∑ºÍ±∞Í∞Ä ÏóÜÏúºÎ©¥ ÎÇ¥Ïö©ÏùÑ Ï±ÑÏö∞ÏßÄ ÎßêÍ≥†, Ìï¥Îãπ ÏßÄÏ†êÏóêÎßå "[ÌôïÏù∏ ÌïÑÏöî: Î¨¥ÏóáÏùÑ ÌôïÏù∏?]" ÎùºÎ≤®ÏùÑ Î∂ôÏù¥ÏÑ∏Ïöî.
3) ÏïîÎ¨µ Ï°∞Í±¥/Ï†ÑÏ†úÎäî raw_textÏóê ÏïîÏãú/ÌëúÌòÑÏù¥ ÏûàÎäî Í≤ΩÏö∞ÏóêÎßå Î™ÖÏãúÏ†ÅÏúºÎ°ú ÌíÄÏñ¥Ïì∞ÏÑ∏Ïöî.
4) ÏïΩÏñ¥/ÎÇ¥Î∂Ä Ïö©Ïñ¥Îäî ÏõêÎ¨∏Ïóê Îì±Ïû•Ìïú Í≤ÉÎßå ÌíÄÏñ¥ ÏÑ§Î™ÖÏùÑ Ï∂îÍ∞ÄÌïòÏÑ∏Ïöî. (ÏõêÎ¨∏Ïóê ÏóÜÏúºÎ©¥ Í∏àÏßÄ)
5) Í∞úÏù∏Ï†ïÎ≥¥(Ï†ÑÌôî/Ïù¥Î©îÏùº/Í≥ÑÏ¢å/ÏÉÅÏÑ∏Ï£ºÏÜå/ÏãùÎ≥ÑÎ≤àÌò∏ Îì±)Îäî ***Î°ú ÎßàÏä§ÌÇπÏùÑ Ïú†ÏßÄÌïòÏÑ∏Ïöî. ÏõêÎ¨∏Ïóê ÏûàÏñ¥ÎèÑ Í∑∏ÎåÄÎ°ú ÎÖ∏Ï∂ú Í∏àÏßÄ.
6) [Q&A Ï†ïÏ±Ö] {qa_policy_text}
7) Í∏∞Ï°¥ section_textÏùò Ï£ºÏ†ú/Î≤îÏúÑÎ•º Î∞îÍæ∏ÏßÄ ÎßàÏÑ∏Ïöî. (Îã§Î•∏ ÏÑπÏÖò Ï£ºÏ†úÎ•º ÏÑûÏñ¥ ÎÑ£ÏßÄ Îßê Í≤É)

[Í∞úÏÑ† Î™©Ìëú]
- ÏïûÎí§ ÏÑπÏÖò ÏóÜÏù¥ÎèÑ Ïù¥ section_textÎßå ÏùΩÍ≥† ÎãµÎ≥Ä Í∞ÄÎä•Ìïú ÏàòÏ§ÄÏúºÎ°ú,
  ÏõêÎ¨∏Ïóê Ìù©Ïñ¥ÏßÑ Í¥ÄÎ†® Í∑úÏπô/Ï°∞Í±¥/Ï±ÑÎÑê/ÏòàÏô∏Î•º Ïù¥ ÏÑπÏÖò ÏïàÏóê ÌÜµÌï©ÌïòÏÑ∏Ïöî.
- Ï§ëÎ≥µ bullet Ï†úÍ±∞, ÌëúÌòÑ Ï†ïÎèà, Ìï≠Î™© Ï†úÎ™©ÏùÑ Î™ÖÌôïÌûà.
- ÎÑàÎ¨¥ Í∏¥ Ìï≠Î™©ÏùÄ Í∞ôÏùÄ Ï£ºÏ†ú ÏïàÏóêÏÑú 2Í∞úÎ°ú Ï™ºÍ∞úÎêò ÌòïÏãùÏùÄ Ïú†ÏßÄÌïòÏÑ∏Ïöî.

[Ï∂úÎ†• ÌòïÏãù(Î∞òÎìúÏãú Ïú†ÏßÄ)]
- Ïò§ÏßÅ Í∞úÏÑ†Îêú section_text Ï†ÑÏ≤¥Î•º plain textÎ°úÎßå Ï∂úÎ†•ÌïòÏÑ∏Ïöî.
- ÏÑπÏÖò ÏãúÏûë: "## "
- Ìï≠Î™© Ï†úÎ™©: "### "
- Î≥∏Î¨∏: "-" bullet
- Q&AÎäî allow_qa=true Ïù¥Í±∞ÎÇò ÏõêÎûò Ï°¥Ïû¨ÌïòÎäî Í≤ΩÏö∞ÏóêÎßå, Ìï≠Î™© ÌïòÎã®Ïóê ÏïÑÎûò ÌòïÏãùÏúºÎ°úÎßå Ìè¨Ìï®:
  - Q: ...
  - A: ...

[Ï∂îÍ∞Ä Í∞ÄÏù¥Îìú]
- raw_textÏóêÏÑú Í∑ºÍ±∞Í∞Ä Î™ÖÌôïÌïú ÎÇ¥Ïö©Îßå 'Î™®ÏïÑÏÑú' ÎÑ£Í≥†, Í∑ºÍ±∞ ÏóÜÎäî Î∂ÄÎ∂ÑÏùÄ Ï±ÑÏö∞ÏßÄ ÏïäÏäµÎãàÎã§.
- "[ÌôïÏù∏ ÌïÑÏöî]"Îäî ÎÇ®Î∞úÌïòÏßÄ ÎßêÍ≥† 'ÌïÑÏàò ÌåêÎã®Ïóê ÌïÑÏöîÌïú ÌïµÏã¨ ÎπàÏπ∏'ÏóêÎßå ÏÇ¨Ïö©ÌïòÏÑ∏Ïöî.
- ÏÑ§Î™Ö/Ìï¥ÏÑ§/JSON Ï∂úÎ†• Í∏àÏßÄ. Ïò§ÏßÅ ÏµúÏ¢Ö ÌÖçÏä§Ìä∏Îßå Ï∂úÎ†•ÌïòÏÑ∏Ïöî.
"""


FINALIZE_SECTION_TEXT_PROMPT_V1 = """ÎãπÏã†ÏùÄ RAG Í≤ÄÏÉâ Ï†ÅÏ§ëÎ•†Í≥º ÎãµÎ≥Ä ÏùºÍ¥ÄÏÑ±ÏùÑ ÎÜíÏù¥Í∏∞ ÏúÑÌï¥ section_text(plain text)Î•º 'ÏµúÏ¢Ö Î¨∏Íµ¨'Î°ú Îã§Îì¨Îäî Ìé∏ÏßëÏûêÏûÖÎãàÎã§.
Ï§ëÏöî: ÏÇ¨Ïã§/Ï†ïÏ±Ö/ÏàòÏπò/Í∏∞Í∞Ñ/Í∏àÏï° Îì± ÏÉàÎ°úÏö¥ ÎÇ¥Ïö©ÏùÑ Ï∂îÍ∞ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî. Ïò§ÏßÅ ÌëúÌòÑÍ≥º Íµ¨Ï°∞Îßå ÏµúÏ†ÅÌôîÌï©ÎãàÎã§.

‚ö†Ô∏è ÎãπÏã†Ïù¥ Îã§Îì¨Ïñ¥Ïïº ÌïòÎäî ÎåÄÏÉÅÏùÄ ÏïÑÎûò [section_text]ÎøêÏûÖÎãàÎã§.
‚ö†Ô∏è [raw_text]Îäî ÏõêÎ≥∏ Ï∞∏Í≥†Ïö©ÏûÖÎãàÎã§. raw_textÏùò Îã§Î•∏ ÏÑπÏÖò ÎÇ¥Ïö©ÏùÑ section_textÏóê Ï∂îÍ∞ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî.
‚ö†Ô∏è ÏûÖÎ†•Îêú section_text Î≤îÏúÑ Î∞ñÏùò ÎÇ¥Ïö©ÏùÄ Ï†àÎåÄ Ìè¨Ìï®ÌïòÏßÄ ÎßàÏÑ∏Ïöî.

[ÏûÖÎ†•]
section_text (Îã§Îì¨ÏùÑ ÎåÄÏÉÅ ‚Äî Ïù¥ ÏÑπÏÖòÎßå Ï∂úÎ†•): {section_text}

raw_text (Ï∞∏Í≥†Ïö© ÏõêÎ≥∏ ‚Äî Ï∂úÎ†•ÌïòÏßÄ ÎßàÏÑ∏Ïöî): {raw_text}

[Ï†àÎåÄ Í∑úÏπô]
1) Ïù¥ ÏÑπÏÖò(section_text)Îßå Îã§Îì¨Ïñ¥ Ï∂úÎ†•. Îã§Î•∏ ÏÑπÏÖò ÎÇ¥Ïö© Ìè¨Ìï® Í∏àÏßÄ
2) ÏÉà Ï†ïÎ≥¥ Ï∂îÍ∞Ä Í∏àÏßÄ(ÏàòÏπò/Í∏∞Í∞Ñ/Í∏àÏï°/Ï†ïÏ±Ö/ÏòàÏô∏/Ï†àÏ∞® Ï∞ΩÏûë Í∏àÏßÄ)
3) ÏõêÎ¨∏/ÌòÑ section_textÏôÄ Îã§Î•∏ ÏÇ¨Ïã§ ÏÉùÏÑ± Í∏àÏßÄ
4) Í∞úÏù∏Ï†ïÎ≥¥ ÎßàÏä§ÌÇπ Ïú†ÏßÄ(***)
5) Í∑ºÍ±∞Í∞Ä Î∂àÎ™ÖÌôïÌïú Î¨∏Ïû• Ï∂îÍ∞Ä Í∏àÏßÄ. Î∂àÎ™ÖÌôïÌïòÎ©¥ "[ÌôïÏù∏ ÌïÑÏöî: ...]"Î•º Ïú†ÏßÄÌïòÍ±∞ÎÇò Îçî Î™ÖÌôïÌûà ÏûëÏÑ±

[ÏµúÏ†ÅÌôî Î™©Ìëú]
- Í≤ÄÏÉâ ÌÇ§ÏõåÎìúÏóê Ïûò Í±∏Î¶¨ÎèÑÎ°ù Ìï≠Î™© Ï†úÎ™©(###)ÏùÑ 'ÏßàÎ¨∏Ìòï ÎòêÎäî ÌÇ§ÏõåÎìúÌòï'ÏúºÎ°ú ÏÑ†Î™ÖÌïòÍ≤å
  Ïòà: "ÌôòÎ∂à Í∑úÏ†ï" ‚Üí "ÌôòÎ∂à Í∑úÏ†ï/Í∏∞Ìïú/ÏúÑÏïΩÍ∏à"
- bulletsÎ•º ÏßßÍ≥† Î≥ëÎ†¨ Íµ¨Ï°∞Î°ú Ï†ïÎ¶¨(Ï§ëÎ≥µ Ï†úÍ±∞)
- Í∞ôÏùÄ ÏùòÎØ∏Ïùò ÌëúÌòÑÏùÑ ÌÜµÏùº(Ïö©Ïñ¥ ÌëúÏ§ÄÌôî)
- Q&AÍ∞Ä Ïù¥ÎØ∏ Ï°¥Ïû¨ÌïòÎ©¥, ÏßàÎ¨∏ÏùÑ Îçî Î™ÖÌôïÌûà ÌïòÎêò ÎãµÏùÄ Í∑∏ÎåÄÎ°ú(ÎÇ¥Ïö© Ï∂îÍ∞Ä Í∏àÏßÄ)
- ÎÑàÎ¨¥ Í∏¥ bulletÏùÄ 2Í∞úÎ°ú Î∂ÑÎ¶¨ÌïòÎêò ÏùòÎØ∏ Ïú†ÏßÄ

[Ï∂úÎ†•]
- Ïò§ÏßÅ Ïù¥ ÏÑπÏÖòÏùò ÏµúÏ¢Ö section_textÎßå plain textÎ°ú Ï∂úÎ†•ÌïòÏÑ∏Ïöî.
- Îã§Î•∏ ÏÑπÏÖò ÎÇ¥Ïö©ÏùÑ Ìï©Ï≥êÏÑú Ï∂úÎ†•ÌïòÏßÄ ÎßàÏÑ∏Ïöî.
- ÌòïÏãù Ïú†ÏßÄ:
  - "## " ÏÑπÏÖò
  - "### " Ìï≠Î™©
  - "-" bullet
  - Q/AÎäî Ï°¥Ïû¨Ìï† ÎïåÎßå Ïú†ÏßÄ
- ÏÑ§Î™Ö/Ìï¥ÏÑ§/JSON Í∏àÏßÄ
"""


@router.post("/doc/{doc_id}/refine-text")
def refine_text(doc_id: str, req: RefineRequest):
    """AI helper: fill (Îß•ÎùΩ Î≥¥Í∞ï), refine (RAG ÏµúÏ†ÅÌôî), recommend (ÌëúÏ§Ä ÌÖúÌîåÎ¶ø Ï†úÏïà)."""
    if not is_llm_available():
        raise HTTPException(status_code=503, detail="LLM ÏÇ¨Ïö© Î∂àÍ∞Ä")

    # Fetch original document raw_text for reference
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    conn.close()

    raw_text = ""
    if doc and doc["raw_text"]:
        raw_text = doc["raw_text"][:4000]

    raw_text_safe = raw_text or "(ÏõêÎ≥∏ Î¨∏ÏÑúÎ•º Ï∞æÏùÑ Ïàò ÏóÜÏäµÎãàÎã§. Í∏∞Ï°¥ ÌÖçÏä§Ìä∏Îßå Ï∞∏Í≥†ÌïòÏÑ∏Ïöî.)"

    try:
        if req.task == "fill":
            allow_qa = _to_bool_allow_qa(req.allow_qa)
            qa_policy_text = (
                "ÏõêÎ¨∏ Í∑ºÍ±∞Í∞Ä Î™ÖÌôïÌïú Í≤ΩÏö∞ÏóêÎßå Q&AÎ•º 1~3Í∞ú Ï∂îÍ∞ÄÌï† Ïàò ÏûàÏäµÎãàÎã§. ÎãµÏù¥ Î∂àÎ™ÖÌôïÌïòÎ©¥ Q&AÎ•º ÎßåÎì§ÏßÄ ÎßêÍ≥† [ÌôïÏù∏ ÌïÑÏöî]Î°ú Ï≤òÎ¶¨ÌïòÏÑ∏Ïöî."
                if allow_qa else
                "Q&AÎäî ÏÉàÎ°ú Ï∂îÍ∞ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî. Í∏∞Ï°¥ Q&AÎßå Ïú†ÏßÄ/Ï†ïÎ¶¨ÌïòÏÑ∏Ïöî."
            )
            prompt = FILL_SECTION_TEXT_PROMPT_V3.format(
                section_text=req.text,
                raw_text=raw_text_safe,
                qa_policy_text=qa_policy_text
            )
        elif req.task == "refine":
            prompt = FINALIZE_SECTION_TEXT_PROMPT_V1.format(
                section_text=req.text,
                raw_text=raw_text_safe
            )
        else:
            # recommend: use fill prompt with Q&A disabled
            prompt = FILL_SECTION_TEXT_PROMPT_V3.format(
                section_text=req.text,
                raw_text=raw_text_safe,
                qa_policy_text="Q&AÎäî ÏÉàÎ°ú Ï∂îÍ∞ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî. Í∏∞Ï°¥ Q&AÎßå Ïú†ÏßÄ/Ï†ïÎ¶¨ÌïòÏÑ∏Ïöî."
            )

        suggestion = call_llm(prompt, temperature=0.3)
        return {"success": True, "suggestion": suggestion}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============ STEP 2: APPROVE ============

@router.post("/doc/{doc_id}/approve")
def approve(doc_id: str):
    """Approve document if no RED issues open and no unresolved evidence failures."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT doc_id FROM documents WHERE doc_id = %s", (doc_id,))
    if not cursor.fetchone():
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")

    cursor.execute("SELECT COUNT(*) as cnt FROM qa_issues WHERE doc_id = %s AND severity = 'RED' AND status = 'OPEN'", (doc_id,))
    red_count = cursor.fetchone()["cnt"]

    if red_count > 0:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Cannot approve: {red_count} RED issues open")

    # Check evidence span indexing failures (non-PII spans with char_start=-1)
    cursor.execute("SELECT section_name, evidence_json FROM manual_sections WHERE doc_id = %s AND evidence_json IS NOT NULL", (doc_id,))
    evidence_failures = []
    for row in cursor.fetchall():
        try:
            spans = json.loads(row["evidence_json"] or "[]")
            for span in spans:
                if span.get("char_start", 0) == -1 and not span.get("is_pii", False):
                    evidence_failures.append(row["section_name"])
                    break
        except (json.JSONDecodeError, TypeError):
            pass

    if evidence_failures:
        conn.close()
        sections_str = ", ".join(evidence_failures[:3])
        raise HTTPException(
            status_code=400,
            detail=f"Evidence span Îß§Ïπ≠ Ïã§Ìå® ÏÑπÏÖò {len(evidence_failures)}Í∞ú ({sections_str}). ManualizeÎ•º Îã§Ïãú Ïã§ÌñâÌïòÏÑ∏Ïöî."
        )

    cursor.execute("UPDATE documents SET status = 'APPROVED', updated_at = %s WHERE doc_id = %s",
                   (datetime.now().isoformat(), doc_id))
    conn.commit()
    conn.close()

    # Auto-reindex after approval
    reindex_result = reindex(doc_id)

    return {"success": True, "doc_id": doc_id, "status": "APPROVED", "reindex": reindex_result}


@router.post("/doc/{doc_id}/publish-all")
def publish_all(doc_id: str):
    """Index PASS sections only into RAG chunks. Update doc status to APPROVED.
    RAG optimize is already done in fill-all, so this just chunks and indexes."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute("SELECT section_name, section_text, gate_status FROM manual_sections WHERE doc_id = %s", (doc_id,))
    all_sections = cursor.fetchall()
    if not all_sections:
        conn.close()
        raise HTTPException(status_code=400, detail="No sections found")

    pass_sections = [s for s in all_sections if s["gate_status"] == "PASS"]
    excluded_count = len(all_sections) - len(pass_sections)

    published_count = 0
    failed_count = 0
    failed_sections = []

    # Delete existing chunks for this doc (will re-insert PASS only)
    cursor.execute("DELETE FROM chunks WHERE doc_id = %s", (doc_id,))

    chunk_size = 500
    overlap = 100

    for sec in pass_sections:
        section_name = sec["section_name"]
        text = sec["section_text"]

        try:
            if not text or text == "Ï†ïÎ≥¥ ÏóÜÏùå":
                continue

            start = 0
            chunk_index = 0
            while start < len(text):
                end = start + chunk_size
                chunk_text = text[start:end]

                if chunk_text.strip():
                    chunk_id = f"chunk_{uuid.uuid4().hex[:8]}"
                    # Generate embedding
                    embedding_blob = None
                    try:
                        emb = get_embedding(chunk_text)
                        embedding_blob = np.array(emb, dtype=np.float32).tobytes()
                    except Exception as emb_err:
                        print(f"[PUBLISH_ALL] Embedding failed for {chunk_id}: {emb_err}")
                    cursor.execute(
                        "INSERT INTO chunks (chunk_id, doc_id, section_name, chunk_index, chunk_text, embedding, created_at) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                        (chunk_id, doc_id, section_name, chunk_index, chunk_text, embedding_blob, datetime.now().isoformat())
                    )
                    chunk_index += 1

                start = end - overlap
                if start >= len(text):
                    break

            published_count += 1
        except Exception as e:
            print(f"[PUBLISH_ALL] Indexing failed for {section_name}: {e}")
            failed_count += 1
            failed_sections.append({"section_name": section_name, "reason": str(e)})

    # Update doc status to APPROVED
    cursor.execute("UPDATE documents SET status = 'APPROVED', updated_at = %s WHERE doc_id = %s",
                   (datetime.now().isoformat(), doc_id))

    conn.commit()
    conn.close()

    # Build FAISS index for this apt_id
    embed_stats = {}
    try:
        conn2 = get_connection()
        c2 = conn2.cursor()
        c2.execute("SELECT apt_id FROM documents WHERE doc_id = %s", (doc_id,))
        doc_row = c2.fetchone()
        conn2.close()
        if doc_row:
            embed_stats = rag_build_index(doc_row["apt_id"])
    except Exception as idx_err:
        print(f"[PUBLISH_ALL] FAISS index build failed: {idx_err}")
        embed_stats = {"error": str(idx_err)}

    published_names = [s["section_name"] for s in pass_sections if s["section_name"] not in [f["section_name"] for f in failed_sections]]
    excluded_names = [s["section_name"] for s in all_sections if s["gate_status"] != "PASS"]
    return {
        "success": True,
        "published_count": published_count,
        "failed_count": failed_count,
        "excluded_count": excluded_count,
        "failed_sections": failed_sections,
        "published_names": published_names,
        "excluded_names": excluded_names,
        "embed_stats": embed_stats
    }


# ============ STEP 2: REINDEX ============

@router.post("/doc/{doc_id}/reindex")
def reindex(doc_id: str):
    """Chunk manual sections for RAG retrieval."""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    sections = cursor.fetchall()
    if not sections:
        conn.close()
        raise HTTPException(status_code=400, detail="No sections to index")
    
    # Delete existing chunks
    cursor.execute("DELETE FROM chunks WHERE doc_id = %s", (doc_id,))
    
    chunk_size = 500
    overlap = 100
    chunk_count = 0
    
    for section in sections:
        text = section["section_text"]
        section_name = section["section_name"]
        
        if not text or text == "Ï†ïÎ≥¥ ÏóÜÏùå":
            continue
        
        # Simple chunking with overlap
        start = 0
        chunk_index = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            if chunk_text.strip():
                chunk_id = f"chunk_{uuid.uuid4().hex[:8]}"
                # Generate embedding
                embedding_blob = None
                try:
                    emb = get_embedding(chunk_text)
                    embedding_blob = np.array(emb, dtype=np.float32).tobytes()
                except Exception as emb_err:
                    print(f"[REINDEX] Embedding failed for {chunk_id}: {emb_err}")
                cursor.execute(
                    "INSERT INTO chunks (chunk_id, doc_id, section_name, chunk_index, chunk_text, embedding, created_at) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                    (chunk_id, doc_id, section_name, chunk_index, chunk_text, embedding_blob, datetime.now().isoformat())
                )
                chunk_count += 1
                chunk_index += 1

            start = end - overlap
            if start >= len(text):
                break

    conn.commit()
    conn.close()

    # Build FAISS index for this apt_id
    embed_stats = {}
    try:
        conn2 = get_connection()
        c2 = conn2.cursor()
        c2.execute("SELECT apt_id FROM documents WHERE doc_id = %s", (doc_id,))
        doc_row = c2.fetchone()
        conn2.close()
        if doc_row:
            embed_stats = rag_build_index(doc_row["apt_id"])
    except Exception as idx_err:
        print(f"[REINDEX] FAISS index build failed: {idx_err}")
        embed_stats = {"error": str(idx_err)}

    return {"success": True, "doc_id": doc_id, "chunk_count": chunk_count, "embed_stats": embed_stats}


# ============ STEP 2: GET SECTIONS ============

@router.get("/doc/{doc_id}/sections")
def get_sections(doc_id: str):
    """Get manual sections for a document (includes chunk-based fields)."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("""SELECT section_name, section_text, gate_status, gate_reasons_json, gate_stale,
                             evidence_json, source_chunk_id, merge_status
                      FROM manual_sections WHERE doc_id = %s""", (doc_id,))
    sections = [dict(row) for row in cursor.fetchall()]
    # Include document-level completed_phases
    cursor.execute("SELECT completed_phases FROM documents WHERE doc_id = %s", (doc_id,))
    doc_row = cursor.fetchone()
    completed_phases = (doc_row["completed_phases"] or "") if doc_row else ""
    conn.close()
    return {"sections": sections, "completed_phases": completed_phases}


def _split_raw_by_headings(raw_text: str) -> list:
    """Split raw_text into chunks by heading patterns."""
    heading_pattern = re.compile(r'^(?:#{1,4}\s+.+|(?:\d+[\.\)]\s*).+|.+[:\uff1a]\s*)$', re.MULTILINE)
    matches = list(heading_pattern.finditer(raw_text))
    if not matches:
        return []
    chunks = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(raw_text)
        heading = m.group().strip().lstrip("#").strip()
        body = raw_text[start:end].strip()
        chunks.append({"heading": heading, "body": body})
    return chunks


def _heading_match(section_name: str, headings: list) -> int:
    """Match section name to heading index by word overlap. Returns -1 if no match."""
    sec_words = set(re.findall(r'[\wÍ∞Ä-Ìû£]+', section_name.lower()))
    if not sec_words:
        return -1
    best_idx, best_score = -1, 0
    for i, h in enumerate(headings):
        h_words = set(re.findall(r'[\wÍ∞Ä-Ìû£]+', h["heading"].lower()))
        overlap = len(sec_words & h_words)
        score = overlap / max(len(sec_words | h_words), 1)
        if score > best_score and score >= 0.3:
            best_score = score
            best_idx = i
    return best_idx


def _body_fallback_match(section_text: str, raw_text: str, min_phrase_len: int = 6) -> str | None:
    """Fallback: find the best matching region in raw_text using keyword phrases from section_text.
    Returns the matched raw_text region or None."""
    if not section_text or not raw_text:
        return None

    # Extract meaningful phrases (lines stripped of markdown markers)
    sec_lines = [ln.lstrip("#>*- \t").strip() for ln in section_text.split("\n")]
    sec_lines = [ln for ln in sec_lines if len(ln) >= min_phrase_len]
    if not sec_lines:
        return None

    # Find matching lines in raw_text
    raw_lines = raw_text.split("\n")
    matched_raw_indices = set()
    for sec_ln in sec_lines:
        # Try exact substring match first
        sec_ln_clean = sec_ln.lstrip("#>*- \t").strip()
        if len(sec_ln_clean) < min_phrase_len:
            continue
        for ri, raw_ln in enumerate(raw_lines):
            raw_ln_clean = raw_ln.lstrip("#>*- \t").strip()
            # Check if significant portion overlaps
            if len(raw_ln_clean) < 4:
                continue
            # Extract keywords (3+ chars) from section line
            sec_words = set(re.findall(r'[\wÍ∞Ä-Ìû£]{3,}', sec_ln_clean.lower()))
            raw_words = set(re.findall(r'[\wÍ∞Ä-Ìû£]{3,}', raw_ln_clean.lower()))
            if not sec_words:
                continue
            overlap = len(sec_words & raw_words) / len(sec_words)
            if overlap >= 0.5:
                matched_raw_indices.add(ri)

    if not matched_raw_indices:
        return None

    # Build contiguous region: from first matched line to last matched line (with 1-line padding)
    first = max(0, min(matched_raw_indices) - 1)
    last = min(len(raw_lines) - 1, max(matched_raw_indices) + 1)
    region = "\n".join(raw_lines[first:last + 1]).strip()
    return region if region else None


@router.get("/doc/{doc_id}/source-map")
def get_source_map(doc_id: str):
    """Get per-section matched raw text.
    Chunk-based: uses doc_chunks JOIN for precise raw_chunk mapping.
    Legacy: falls back to heading matching."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT raw_text FROM documents WHERE doc_id = %s", (doc_id,))
    doc = cursor.fetchone()
    if not doc:
        conn.close()
        raise HTTPException(status_code=404, detail="Document not found")

    cursor.execute("SELECT section_name, source_chunk_id, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    sections = cursor.fetchall()
    conn.close()

    raw_text = doc["raw_text"] or ""

    # Check if any section has source_chunk_id (chunk-based document)
    has_chunks = any(row["source_chunk_id"] for row in sections)

    if has_chunks:
        # Chunk-based: JOIN doc_chunks for raw_chunk
        conn2 = get_connection()
        cursor2 = conn2.cursor()
        source_map = {}
        matched_count = 0
        for row in sections:
            name = row["section_name"]
            chunk_id = row["source_chunk_id"]
            if chunk_id:
                cursor2.execute("SELECT raw_chunk FROM doc_chunks WHERE chunk_id = %s", (chunk_id,))
                chunk_row = cursor2.fetchone()
                if chunk_row and chunk_row["raw_chunk"]:
                    source_map[name] = chunk_row["raw_chunk"]
                    matched_count += 1
                else:
                    source_map[name] = None
            else:
                source_map[name] = None
        conn2.close()
        # Body fallback for unmatched sections
        for row in sections:
            name = row["section_name"]
            if source_map.get(name) is None:
                fallback = _body_fallback_match(row["section_text"] or "", raw_text)
                if fallback:
                    source_map[name] = fallback
                    matched_count += 1
        return {"source_map": source_map, "raw_text": raw_text, "matched": matched_count}

    # Legacy: heading-based matching
    section_names = [row["section_name"] for row in sections]
    section_texts = {row["section_name"]: row["section_text"] or "" for row in sections}
    chunks = _split_raw_by_headings(raw_text)
    if not chunks:
        # No headings found ‚Äî try body fallback for every section
        source_map = {}
        for name in section_names:
            source_map[name] = _body_fallback_match(section_texts[name], raw_text)
        return {"source_map": source_map, "raw_text": raw_text, "matched": sum(1 for v in source_map.values() if v is not None)}

    match_indices = []
    for sec_name in section_names:
        idx = _heading_match(sec_name, chunks)
        match_indices.append(idx)

    source_map = {}
    for i, sec_name in enumerate(section_names):
        idx = match_indices[i]
        if idx < 0:
            # Heading match failed ‚Äî try body fallback
            source_map[sec_name] = _body_fallback_match(section_texts[sec_name], raw_text)
            continue
        next_chunk_idx = len(chunks)
        for j in range(i + 1, len(section_names)):
            if match_indices[j] > idx:
                next_chunk_idx = match_indices[j]
                break
        parts = [chunks[k]["body"] for k in range(idx, min(next_chunk_idx, len(chunks)))]
        source_map[sec_name] = "\n\n".join(parts)

    return {"source_map": source_map, "raw_text": raw_text, "matched": sum(1 for v in source_map.values() if v is not None)}


@router.get("/doc/{doc_id}/issues")
def get_issues(doc_id: str):
    """Get QA issues for a document."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT issue_id, severity, issue_type, message, suggestion, status FROM qa_issues WHERE doc_id = %s", (doc_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return issues


# ============ STEP 3: CHAT WITH RAG ============

class ChatRequest(BaseModel):
    apt_id: str
    client_id: str = "default"
    conversation_id: Optional[str] = None
    message: str
    model: Optional[str] = None


CHAT_PROMPT = """ÎãπÏã†ÏùÄ ÏïÑÌååÌä∏ Ïª§ÎÆ§ÎãàÌã∞ Í∑úÏ†ïÏóê ÎåÄÌï¥ ÎãµÎ≥ÄÌïòÎäî AI Ïñ¥ÏãúÏä§ÌÑ¥Ìä∏ÏûÖÎãàÎã§.

ÏïÑÎûò Î¨∏ÏÑú ÎÇ¥Ïö©ÏùÑ Î∞îÌÉïÏúºÎ°ú ÏÇ¨Ïö©Ïûê ÏßàÎ¨∏Ïóê ÎãµÎ≥ÄÌïòÏÑ∏Ïöî.
Î¨∏ÏÑúÏóê ÏóÜÎäî ÎÇ¥Ïö©ÏùÄ ÎãµÎ≥ÄÌïòÏßÄ ÎßàÏÑ∏Ïöî. Î∞òÎìúÏãú ÏïÑÎûò JSON ÌòïÏãùÏúºÎ°úÎßå ÏùëÎãµÌïòÏÑ∏Ïöî.

Î¨∏ÏÑú ÎÇ¥Ïö©:
{context}

ÏÇ¨Ïö©Ïûê ÏßàÎ¨∏: {question}

JSON ÌòïÏãù:
{{
  "reply_text": "ÎãµÎ≥Ä ÎÇ¥Ïö© (Í∞ÑÍ≤∞ÌïòÍ≤å)",
  "citations": [
    {{"doc_id": "Î¨∏ÏÑúID", "doc_title": "Î¨∏ÏÑúÎ™Ö", "section_name": "ÏÑπÏÖòÎ™Ö", "snippet": "Ïù∏Ïö© Î∂ÄÎ∂Ñ 120Ïûê Ïù¥ÎÇ¥"}}
  ],
  "confidence": "HIGH|MED|LOW",
  "next_question": null ÎòêÎäî "Ï∂îÍ∞Ä ÏßàÎ¨∏(Í∑ºÍ±∞ Î∂ÄÏ°±Ïãú)",
  "actions": []
}}

Í∑úÏπô:
- citationsÏù¥ ÏóÜÏúºÎ©¥ confidenceÎäî LOWÎ°ú ÏÑ§Ï†ï
- LOWÏùº Í≤ΩÏö∞ next_questionÏóê ÌôïÏù∏Ìï† ÏßàÎ¨∏ 1Í∞ú Ìè¨Ìï®
- ÏòàÏïΩÏÉùÏÑ±/Î¨∏ÏûêÎ∞úÏÜ° Îì± Ïã§Ìñâ Ïï°ÏÖòÏùÄ Í∏àÏßÄ (actionsÎäî Ìï≠ÏÉÅ [])
- Í∑ºÍ±∞ ÏóÜÏù¥ Ï∂îÏ∏°ÌïòÏßÄ ÎßàÏÑ∏Ïöî"""


def keyword_search(query: str, chunks: list, top_k: int = 5) -> list:
    """Simple keyword-based search scoring."""
    keywords = set(query.lower().replace("?", "").replace(".", "").split())
    scored = []
    for chunk in chunks:
        text = chunk["chunk_text"].lower()
        score = sum(1 for kw in keywords if kw in text)
        if score > 0:
            scored.append((score, chunk))
    scored.sort(key=lambda x: -x[0])
    return [c for _, c in scored[:top_k]]


@router.post("/chat")
def chat(req: ChatRequest):
    """Chat with RAG retrieval and citations."""
    import time as _time
    t_start = _time.time()
    timing = {}

    conn = get_connection()
    cursor = conn.cursor()

    # Get or create conversation
    if req.conversation_id:
        conv_id = req.conversation_id
        cursor.execute("UPDATE conversations SET last_at = %s WHERE conversation_id = %s",
                       (datetime.now().isoformat(), conv_id))
    else:
        conv_id = f"conv_{uuid.uuid4().hex[:12]}"
        cursor.execute(
            "INSERT INTO conversations (conversation_id, apt_id, client_id, created_at, last_at) VALUES (%s, %s, %s, %s, %s)",
            (conv_id, req.apt_id, req.client_id, datetime.now().isoformat(), datetime.now().isoformat())
        )
    
    # Save user message
    user_msg_id = f"msg_{uuid.uuid4().hex[:8]}"
    cursor.execute(
        "INSERT INTO messages (msg_id, conversation_id, role, text, created_at) VALUES (%s, %s, 'user', %s, %s)",
        (user_msg_id, conv_id, req.message, datetime.now().isoformat())
    )
    
    # FAISS vector search with keyword fallback
    t_rag = _time.time()
    search_method = "faiss"
    top_chunks = []
    try:
        top_chunks = faiss_search(req.apt_id, req.message, top_k=5)
    except Exception as faiss_err:
        print(f"[CHAT] FAISS search failed, falling back to keyword: {faiss_err}")

    if not top_chunks:
        # Fallback: keyword search
        search_method = "keyword"
        cursor.execute("""
            SELECT c.chunk_id, c.doc_id, c.section_name, c.chunk_text, d.title as doc_title
            FROM chunks c
            JOIN documents d ON c.doc_id = d.doc_id
            WHERE d.apt_id = %s AND d.status = 'APPROVED'
        """, (req.apt_id,))
        all_chunks = [dict(row) for row in cursor.fetchall()]
        top_chunks = keyword_search(req.message, all_chunks, top_k=5)
    timing["rag_s"] = round(_time.time() - t_rag, 2)
    
    # Default response
    response = {
        "conversation_id": conv_id,
        "reply_text": "",
        "citations": [],
        "confidence": "LOW",
        "next_question": None,
        "actions": []
    }
    
    if not top_chunks:
        response["reply_text"] = "Ï£ÑÏÜ°Ìï©ÎãàÎã§. Í¥ÄÎ†® Ï†ïÎ≥¥Î•º Ï∞æÏùÑ Ïàò ÏóÜÏäµÎãàÎã§."
        response["next_question"] = "Ïñ¥Îñ§ ÎÇ¥Ïö©Ïóê ÎåÄÌï¥ Îçî ÏïåÍ≥† Ïã∂ÏúºÏã†Í∞ÄÏöî?"
    else:
        # Build context
        context_parts = []
        for chunk in top_chunks:
            context_parts.append(f"[Î¨∏ÏÑú: {chunk['doc_title']} / ÏÑπÏÖò: {chunk['section_name']}]\n{chunk['chunk_text']}")
        context = "\n\n".join(context_parts)
        
        if is_llm_available():
            try:
                t_llm = _time.time()
                content = call_llm(CHAT_PROMPT.format(context=context, question=req.message), temperature=0.3, model=req.model)
                timing["llm_s"] = round(_time.time() - t_llm, 2)
                if content:
                    json_match = re.search(r'\{[\s\S]*\}', content)
                    if json_match:
                        parsed = _clean_llm_json(json_match.group())
                        response["reply_text"] = parsed.get("reply_text", "")
                        response["citations"] = parsed.get("citations", [])
                        response["confidence"] = parsed.get("confidence", "MED")
                        response["next_question"] = parsed.get("next_question")
            except Exception as e:
                response["reply_text"] = f"LLM Ïò§Î•ò: {str(e)[:50]}"
        
        if not response["reply_text"]:
            # Mock response without LLM
            response["reply_text"] = f"Î¨∏ÏÑúÏóêÏÑú {len(top_chunks)}Í∞úÏùò Í¥ÄÎ†® Ï†ïÎ≥¥Î•º Ï∞æÏïòÏäµÎãàÎã§."
            response["citations"] = [
                {
                    "doc_id": c["doc_id"],
                    "doc_title": c["doc_title"],
                    "section_name": c["section_name"],
                    "snippet": c["chunk_text"][:150] + "..." if len(c["chunk_text"]) > 150 else c["chunk_text"]
                } for c in top_chunks[:2]
            ]
            response["confidence"] = "MED" if top_chunks else "LOW"
    
    # Save assistant message
    asst_msg_id = f"msg_{uuid.uuid4().hex[:8]}"
    meta_json = json.dumps({
        "citations": response["citations"],
        "confidence": response["confidence"],
        "retrieval_count": len(top_chunks)
    }, ensure_ascii=False)
    cursor.execute(
        "INSERT INTO messages (msg_id, conversation_id, role, text, meta_json, created_at) VALUES (%s, %s, 'assistant', %s, %s, %s)",
        (asst_msg_id, conv_id, response["reply_text"], meta_json, datetime.now().isoformat())
    )
    
    conn.commit()
    conn.close()

    timing["total_s"] = round(_time.time() - t_start, 2)
    response["timing"] = timing
    response["search_method"] = search_method

    return response


# ============ STEP 3: IMPROVEMENTS GENERATOR ============

@router.post("/improvements/generate")
def generate_improvements(apt_id: str):
    """Generate improvement suggestions from chat logs."""
    conn = get_connection()
    cursor = conn.cursor()
    
    # Get recent assistant messages with LOW confidence or empty citations
    cursor.execute("""
        SELECT m.msg_id, m.text, m.meta_json, m.created_at, c.apt_id
        FROM messages m
        JOIN conversations c ON m.conversation_id = c.conversation_id
        WHERE c.apt_id = %s AND m.role = 'assistant'
        ORDER BY m.created_at DESC
        LIMIT 50
    """, (apt_id,))
    messages = cursor.fetchall()
    
    suggestions = []
    seen_topics = set()
    
    for msg in messages:
        meta = json.loads(msg["meta_json"] or "{}")
        confidence = meta.get("confidence", "HIGH")
        citations = meta.get("citations", [])
        
        # Check for LOW confidence or empty citations
        if confidence == "LOW" or not citations:
            # Extract topic from message text
            text = msg["text"][:100]
            topic_key = text[:30]
            
            if topic_key not in seen_topics and len(suggestions) < 5:
                seen_topics.add(topic_key)
                
                # Determine target section
                target_section = "ÏòàÏô∏/Î¨∏Ïùò/Í∂åÌïú"
                if "ÌôòÎ∂à" in text or "Ï†ïÏÇ∞" in text:
                    target_section = "ÌôòÎ∂à/ÏúÑÏïΩ/Ï†ïÏÇ∞"
                elif "ÏòàÏïΩ" in text or "Ï∑®ÏÜå" in text:
                    target_section = "ÏòàÏïΩ/Ï∑®ÏÜå/Î≥ÄÍ≤Ω"
                elif "Ïö¥ÏòÅ" in text or "ÏãúÍ∞Ñ" in text:
                    target_section = "Ïö¥ÏòÅÏãúÍ∞Ñ/Ìú¥Î¨¥"
                
                suggestions.append({
                    "title": f"Ï†ïÎ≥¥ Î≥¥ÏôÑ ÌïÑÏöî: {text[:30]}...",
                    "reason": f"confidence={confidence}, citations={len(citations)}Í∞ú",
                    "target_section_name": target_section,
                    "proposed_patch": ""
                })
    
    # Get the latest APPROVED doc for this apt
    cursor.execute(
        "SELECT doc_id FROM documents WHERE apt_id = %s AND status = 'APPROVED' ORDER BY version DESC LIMIT 1",
        (apt_id,)
    )
    doc_row = cursor.fetchone()
    target_doc_id = doc_row["doc_id"] if doc_row else None
    
    # Save suggestions
    for sug in suggestions:
        sug_id = f"sug_{uuid.uuid4().hex[:8]}"
        cursor.execute("""
            INSERT INTO improve_suggestions (sug_id, apt_id, title, reason, proposed_patch, target_doc_id, target_section_name, status, created_at, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'PENDING', %s, %s)
        """, (sug_id, apt_id, sug["title"], sug["reason"], sug["proposed_patch"], target_doc_id, sug["target_section_name"],
              datetime.now().isoformat(), datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "count": len(suggestions), "suggestions": suggestions}


# ============ STEP 3: ONE-CLICK PATCH APPLY ============

PATCH_PROMPT = """Î¨∏ÏÑú ÏÑπÏÖòÏóê Ï∂îÍ∞ÄÌï† FAQ Ìï≠Î™©ÏùÑ ÏÉùÏÑ±ÌïòÏÑ∏Ïöî.

Ï†úÏïà Ï†úÎ™©: {title}
Ï†úÏïà Ïù¥Ïú†: {reason}
ÎåÄÏÉÅ ÏÑπÏÖò: {section_name}

ÌòÑÏû¨ ÏÑπÏÖò ÎÇ¥Ïö©:
{section_text}

Í∑úÏπô:
- ÏóÜÎäî Í∑úÏ†ïÏùÑ ÎßåÎì§ÏßÄ ÎßàÏÑ∏Ïöî
- ÌôïÏã§ÌïòÏßÄ ÏïäÏúºÎ©¥ "ÌôïÏù∏ ÌïÑÏöî" ÌòïÌÉúÎ°ú ÏûëÏÑ±
- Í∞ÑÍ≤∞Ìïú Q&A ÌòïÏãùÏúºÎ°ú ÏûëÏÑ±

Ï∂úÎ†• ÌòïÏãù (Ï∂îÍ∞ÄÌï† ÌÖçÏä§Ìä∏Îßå):
---FAQ---
Q: ÏßàÎ¨∏
A: ÎãµÎ≥Ä
---"""


@router.post("/improvements/{sug_id}/apply")
def apply_improvement(sug_id: str):
    """Apply improvement suggestion with one-click patch."""
    conn = get_connection()
    cursor = conn.cursor()
    
    # Get suggestion
    cursor.execute("SELECT * FROM improve_suggestions WHERE sug_id = %s", (sug_id,))
    sug = cursor.fetchone()
    if not sug:
        conn.close()
        raise HTTPException(status_code=404, detail="Suggestion not found")
    
    if sug["status"] == "APPLIED":
        conn.close()
        raise HTTPException(status_code=400, detail="Already applied")
    
    target_doc_id = sug["target_doc_id"]
    target_section = sug["target_section_name"]
    
    # Get current section text
    cursor.execute(
        "SELECT section_id, section_text FROM manual_sections WHERE doc_id = %s AND section_name = %s",
        (target_doc_id, target_section)
    )
    section_row = cursor.fetchone()
    
    if not section_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Target section not found")
    
    current_text = section_row["section_text"]
    section_id = section_row["section_id"]
    
    # Generate patch
    patch_text = ""
    if is_llm_available():
        try:
            prompt = PATCH_PROMPT.format(
                title=sug["title"],
                reason=sug["reason"],
                section_name=target_section,
                section_text=current_text[:2000]
            )
            patch_text = call_llm(prompt, temperature=0.3)
        except Exception as e:
            print(f"[PATCH] Error: {e}")
            patch_text = f"\n\n---FAQ---\nQ: {sug['title']}\nA: ÌôïÏù∏ ÌïÑÏöî - Í¥ÄÎ¶¨ÏûêÏóêÍ≤å Î¨∏ÏùòÌïòÏÑ∏Ïöî."
    
    if not patch_text:
        patch_text = f"\n\n---FAQ---\nQ: {sug['title']}\nA: ÌôïÏù∏ ÌïÑÏöî - Í¥ÄÎ¶¨ÏûêÏóêÍ≤å Î¨∏ÏùòÌïòÏÑ∏Ïöî."
    
    # Append patch to section
    new_text = current_text + "\n" + patch_text
    cursor.execute(
        "UPDATE manual_sections SET section_text = %s WHERE section_id = %s",
        (new_text, section_id)
    )
    
    # Update suggestion status
    cursor.execute(
        "UPDATE improve_suggestions SET status = 'APPLIED', updated_at = %s WHERE sug_id = %s",
        (datetime.now().isoformat(), sug_id)
    )
    
    conn.commit()
    conn.close()
    
    # Reindex the document
    reindex_result = reindex(target_doc_id)
    
    return {"success": True, "sug_id": sug_id, "status": "APPLIED", "reindexed": reindex_result}


# ============ STEP 3: API SPEC EXTRACTOR ============

API_SPEC_PROMPT = """ÏïÑÎûò Îß§Îâ¥Ïñº ÏÑπÏÖòÏóêÏÑú ÏãúÏä§ÌÖú APIÍ∞Ä ÌïÑÏöîÌïú ÏùòÎèÑ(intent)Î•º Ï∂îÏ∂úÌïòÏÑ∏Ïöî.

Îß§Îâ¥Ïñº ÎÇ¥Ïö©:
{sections_text}

ÌíàÏßà Ïù¥Ïäà (API_NEEDED):
{api_issues}

Í∞Å intentÏóê ÎåÄÌï¥ API Ïä§ÌéôÏùÑ JSON Î∞∞Ïó¥Î°ú Î∞òÌôòÌïòÏÑ∏Ïöî:
[
  {{
    "intent_name": "ÏòàÏïΩ ÏÉùÏÑ±",
    "endpoint": "/api/booking/create",
    "method": "POST",
    "request_fields": ["member_id", "class_id", "date"],
    "response_fields": ["booking_id", "status"],
    "auth": "ÏûÖÏ£ºÎØº|Í¥ÄÎ¶¨Ïûê|ÏãúÏä§ÌÖú",
    "notes": "ÎπÑÍ≥†"
  }}
]

JSON Î∞∞Ïó¥Îßå Î∞òÌôòÌïòÏÑ∏Ïöî."""


@router.post("/doc/{doc_id}/extract-api-spec")
def extract_api_spec(doc_id: str):
    """Extract API specifications from document."""
    conn = get_connection()
    cursor = conn.cursor()
    
    # Get sections
    cursor.execute("SELECT section_name, section_text FROM manual_sections WHERE doc_id = %s", (doc_id,))
    sections = cursor.fetchall()
    
    # Get API_NEEDED issues
    cursor.execute("SELECT message FROM qa_issues WHERE doc_id = %s AND issue_type = 'API_NEEDED'", (doc_id,))
    api_issues = [row["message"] for row in cursor.fetchall()]
    
    sections_text = "\n\n".join([f"[{s['section_name']}]\n{s['section_text']}" for s in sections])
    
    specs = []
    
    if is_llm_available():
        try:
            prompt = API_SPEC_PROMPT.format(
                sections_text=sections_text[:4000],
                api_issues=", ".join(api_issues)
            )
            content = call_llm(prompt, temperature=0.3)
            if content:
                json_match = re.search(r'\[[\s\S]*\]', content)
                if json_match:
                    specs = _clean_llm_json(json_match.group())
        except Exception as e:
            print(f"[API_SPEC] Error: {e}")
    
    # Fallback: generate from API_NEEDED issues
    if not specs:
        for issue in api_issues:
            intent = issue.replace("ÏãúÏä§ÌÖú Ïó∞Îèô ÌïÑÏöî: ", "").replace("'", "")
            specs.append({
                "intent_name": intent,
                "endpoint": f"/api/{intent.replace(' ', '-').lower()}",
                "method": "POST",
                "request_fields": ["member_id"],
                "response_fields": ["status", "message"],
                "auth": "Í¥ÄÎ¶¨Ïûê",
                "notes": "ÏûêÎèô Ï∂îÏ∂úÎê® - Í≤ÄÌÜ† ÌïÑÏöî"
            })
    
    # Clear and save specs
    cursor.execute("DELETE FROM api_specs WHERE doc_id = %s", (doc_id,))
    for spec in specs:
        spec_id = f"spec_{uuid.uuid4().hex[:8]}"
        cursor.execute("""
            INSERT INTO api_specs (spec_id, doc_id, intent, endpoint, method, req_fields_json, resp_fields_json, auth, errors_json, created_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (spec_id, doc_id, spec.get("intent_name", ""), spec.get("endpoint", ""), spec.get("method", "POST"),
              json.dumps(spec.get("request_fields", []), ensure_ascii=False),
              json.dumps(spec.get("response_fields", []), ensure_ascii=False),
              spec.get("auth", ""), "[]", datetime.now().isoformat()))
    
    conn.commit()
    conn.close()
    
    return {"success": True, "doc_id": doc_id, "spec_count": len(specs), "specs": specs}


@router.get("/doc/{doc_id}/api-spec/export")
def export_api_spec(doc_id: str, format: str = "json"):
    """Export API specs as JSON or YAML."""
    conn = get_connection()
    cursor = conn.cursor()
    
    cursor.execute("""
        SELECT intent, endpoint, method, req_fields_json, resp_fields_json, auth, errors_json
        FROM api_specs WHERE doc_id = %s
    """, (doc_id,))
    rows = cursor.fetchall()
    conn.close()
    
    specs = []
    for row in rows:
        specs.append({
            "intent": row["intent"],
            "endpoint": row["endpoint"],
            "method": row["method"],
            "request_fields": json.loads(row["req_fields_json"] or "[]"),
            "response_fields": json.loads(row["resp_fields_json"] or "[]"),
            "auth": row["auth"],
            "errors": json.loads(row["errors_json"] or "[]")
        })
    
    if format == "yaml":
        # Simple YAML conversion
        yaml_lines = ["api_specs:"]
        for spec in specs:
            yaml_lines.append(f"  - intent: {spec['intent']}")
            yaml_lines.append(f"    endpoint: {spec['endpoint']}")
            yaml_lines.append(f"    method: {spec['method']}")
            yaml_lines.append(f"    auth: {spec['auth']}")
            yaml_lines.append(f"    request_fields: {spec['request_fields']}")
            yaml_lines.append(f"    response_fields: {spec['response_fields']}")
        return {"format": "yaml", "content": "\n".join(yaml_lines)}
    
    return {"format": "json", "specs": specs}


# ============ STEP 3: GET SUGGESTIONS ============

@router.get("/improvements")
def list_improvements(apt_id: Optional[str] = None):
    """List improvement suggestions."""
    conn = get_connection()
    cursor = conn.cursor()
    
    if apt_id:
        cursor.execute("""
            SELECT s.*, a.name as apt_name
            FROM improve_suggestions s
            LEFT JOIN apartments a ON s.apt_id = a.apt_id
            WHERE s.apt_id = %s
            ORDER BY s.created_at DESC
        """, (apt_id,))
    else:
        cursor.execute("""
            SELECT s.*, a.name as apt_name
            FROM improve_suggestions s
            LEFT JOIN apartments a ON s.apt_id = a.apt_id
            ORDER BY s.created_at DESC
        """)
    
    suggestions = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return suggestions


@router.get("/conversations")
def list_conversations(apt_id: Optional[str] = None):
    """List conversations."""
    conn = get_connection()
    cursor = conn.cursor()
    
    if apt_id:
        cursor.execute("SELECT * FROM conversations WHERE apt_id = %s ORDER BY last_at DESC", (apt_id,))
    else:
        cursor.execute("SELECT * FROM conversations ORDER BY last_at DESC")
    
    convs = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return convs


@router.get("/conversations/{conversation_id}/messages")
def get_conversation_messages(conversation_id: str):
    """Get messages for a conversation."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM messages WHERE conversation_id = %s ORDER BY created_at", (conversation_id,))
    messages = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return messages


# ============ STEP 4-1: BRANCH CLASS SYNC ============

@router.post("/branch/{branch_id}/classes/sync")
def sync_branch_classes(
    branch_id: str,
    apt_id: str = Form(...),
    classes_json: str = Form(...),
    asof: Optional[str] = Form(None)
):
    """Sync branch classes from form-urlencoded data.
    
    Accepts application/x-www-form-urlencoded with:
    - apt_id: apartment ID (required)
    - classes_json: JSON string array of class objects (required)
    - asof: timestamp string (optional)
    """
    # Parse classes_json
    try:
        classes = json.loads(classes_json)
        if not isinstance(classes, list):
            raise ValueError("classes_json must be a JSON array")
    except json.JSONDecodeError as e:
        return {
            "success": False,
            "error": str(e),
            "preview": classes_json[:100] if classes_json else "",
            "hint": "classes_json must be valid JSON array"
        }
    except ValueError as e:
        return {
            "success": False,
            "error": str(e),
            "preview": classes_json[:100] if classes_json else "",
            "hint": "classes_json must be valid JSON array"
        }
    
    conn = get_connection()
    cursor = conn.cursor()
    
    now = datetime.now().isoformat()
    asof_value = asof or now
    upserted = 0
    
    for cls in classes:
        class_id = cls.get("class_id") or cls.get("id") or f"cls_{uuid.uuid4().hex[:8]}"
        name = cls.get("name", "")
        start = cls.get("start", "")
        end = cls.get("end", "")
        capacity = cls.get("capacity", 0)
        reserved = cls.get("reserved", 0)
        
        # Upsert using INSERT OR REPLACE
        cursor.execute("""
            REPLACE INTO branch_class_cache
            (apt_id, branch_id, class_id, name, start, end, capacity, reserved, asof, updated_at)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (apt_id, branch_id, class_id, name, start, end, capacity, reserved, asof_value, now))
        upserted += 1
    
    conn.commit()
    conn.close()
    
    return {"success": True, "upserted": upserted, "branch_id": branch_id, "apt_id": apt_id}


@router.get("/branch/{branch_id}/classes")
def get_branch_classes(branch_id: str, apt_id: str, date: Optional[str] = None):
    """Get cached classes for a branch."""
    conn = get_connection()
    cursor = conn.cursor()
    
    if date:
        # Filter by date (classes where start contains the date string)
        cursor.execute("""
            SELECT * FROM branch_class_cache
            WHERE branch_id = %s AND apt_id = %s AND start LIKE %s
            ORDER BY start
        """, (branch_id, apt_id, f"{date}%"))
    else:
        cursor.execute("""
            SELECT * FROM branch_class_cache
            WHERE branch_id = %s AND apt_id = %s
            ORDER BY start
        """, (branch_id, apt_id))
    
    classes = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return {"branch_id": branch_id, "apt_id": apt_id, "classes": classes, "count": len(classes)}


@router.get("/classes")
def list_all_classes(apt_id: Optional[str] = None):
    """List all cached classes, optionally filtered by apt_id."""
    conn = get_connection()
    cursor = conn.cursor()
    
    if apt_id:
        cursor.execute("SELECT * FROM branch_class_cache WHERE apt_id = %s ORDER BY start", (apt_id,))
    else:
        cursor.execute("SELECT * FROM branch_class_cache ORDER BY start")
    
    classes = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    return {"classes": classes, "count": len(classes)}
